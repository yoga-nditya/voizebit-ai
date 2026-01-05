import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config_new import FILES_DIR

from .helpers import (
    format_tanggal_indonesia,
    replace_everywhere_keep_format,
)


def set_run_font(run, font_name="Times New Roman", size=10, bold=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:cs'), font_name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def style_cell_paragraph(cell, align="left", left_indent_pt=0, font="Times New Roman", size=10):
    if not cell.paragraphs:
        cell.add_paragraph("")
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if left_indent_pt and align == "left":
        p.paragraph_format.left_indent = Pt(left_indent_pt)
    for r in p.runs:
        set_run_font(r, font, size)


# =========================
# Create MoU DOCX (PIHAK KETIGA: hanya perusahaan + alamat; nama/jabatan ttd TIDAK disentuh)
# =========================

def create_mou_docx(mou_data: dict, fname_base: str) -> str:
    template_path = "tamplate MoU.docx"
    if not os.path.exists(template_path):
        raise Exception("Template MoU tidak ditemukan. Pastikan file 'tamplate MoU.docx' ada di root project.")

    doc = Document(template_path)

    pihak1 = (mou_data.get("pihak_pertama") or "").strip()
    pihak2 = (mou_data.get("pihak_kedua") or "").strip()
    pihak3 = (mou_data.get("pihak_ketiga") or "").strip()

    alamat1 = (mou_data.get("alamat_pihak_pertama") or "").strip()
    alamat3 = (mou_data.get("alamat_pihak_ketiga") or "").strip()

    ttd1 = (mou_data.get("ttd_pihak_pertama") or "").strip()
    jab1 = (mou_data.get("jabatan_pihak_pertama") or "").strip()

    nomor_full = (mou_data.get("nomor_surat") or "").strip()
    tanggal_text = format_tanggal_indonesia(datetime.now())

    contoh_pihak1_candidates = [
        "PT. PANPAN LUCKY INDONESIA",
        "PT. Panpan Lucky Indonesia",
        "PT PANPAN LUCKY INDONESIA",
        "PT Panpan Lucky Indonesia",
    ]
    contoh_pihak2_candidates = [
        "PT. SARANA TRANS BERSAMA JAYA",
        "PT Sarana Trans Bersama Jaya",
        "PT SARANA TRANS BERSAMA JAYA",
        "PT Sarana Trans Bersama Jaya",
    ]
    contoh_pihak3_candidates = [
        "PT. HARAPAN BARU SEJAHTERA PLASTIK",
        "PT Harapan Baru Sejahtera Plastik",
        "PT HARAPAN BARU SEJAHTERA PLASTIK",
        "PT Harapan Baru Sejahtera Plastik",
    ]

    replace_everywhere_keep_format(doc, contoh_pihak1_candidates, pihak1.upper() if pihak1 else "")
    replace_everywhere_keep_format(doc, contoh_pihak2_candidates, pihak2.upper() if pihak2 else "")
    replace_everywhere_keep_format(doc, contoh_pihak3_candidates, pihak3.upper() if pihak3 else "")

    def replace_no_line(container_paragraphs):
        for p in container_paragraphs:
            if re.search(r'\bNo\s*:', p.text, flags=re.IGNORECASE):
                for run in p.runs:
                    if re.search(r'\bNo\s*:', run.text, flags=re.IGNORECASE):
                        run.text = re.sub(r'\bNo\s*:\s*.*', f"No : {nomor_full}", run.text, flags=re.IGNORECASE)
                        return True
                if p.runs:
                    p.runs[0].text = f"No : {nomor_full}"
                    for r in p.runs[1:]:
                        r.text = ""
                    return True
        return False

    replace_no_line(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if replace_no_line(cell.paragraphs):
                    break

    kalimat_tanggal = f"Pada hari ini {tanggal_text} kami yang bertanda tangan di bawah ini :"
    def replace_pada_hari_ini(container_paragraphs):
        for p in container_paragraphs:
            if "Pada hari ini" in p.text and "bertanda tangan" in p.text:
                if p.runs:
                    p.runs[0].text = kalimat_tanggal
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.add_run(kalimat_tanggal)
                return True
        return False

    replace_pada_hari_ini(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if replace_pada_hari_ini(cell.paragraphs):
                    break

    if ttd1:
        replace_everywhere_keep_format(doc, ["Huang Feifang"], ttd1)
    if jab1:
        replace_everywhere_keep_format(doc, ["Direktur Utama"], jab1)

    contoh_alamat_p1_candidates = [
        "Jl. Raya Serang KM. 22 No. 30, Desa Pasir Bolang, Kec Tigaraksa, Tangerang Banten",
        "Jl. Raya Serang KM. 22 No. 30, Desa Pasir Bolang, Kec. Tigaraksa, Tangerang Banten",
    ]
    if alamat1:
        replace_everywhere_keep_format(doc, contoh_alamat_p1_candidates, alamat1)

    contoh_alamat_p3_candidates = [
        "Jl. Karawang – Bekasi KM. 1 Bojongsari, Kec. Kedungwaringin, Kab. Bekasi – Jawa Barat",
        "Jl. Karawang - Bekasi KM. 1 Bojongsari, Kec. Kedungwaringin, Kab. Bekasi - Jawa Barat",
    ]
    if alamat3:
        replace_everywhere_keep_format(doc, contoh_alamat_p3_candidates, alamat3)

    items = mou_data.get("items_limbah") or []
    target_table = None
    for t in doc.tables:
        if not t.rows:
            continue
        header_text = " ".join([c.text.strip() for c in t.rows[0].cells])
        if ("Jenis Limbah" in header_text) and ("Kode Limbah" in header_text):
            target_table = t
            break

    if target_table is not None:
        while len(target_table.rows) > 1:
            target_table._tbl.remove(target_table.rows[1]._tr)
        for i, it in enumerate(items, start=1):
            row = target_table.add_row()
            cells = row.cells
            if len(cells) >= 1:
                cells[0].text = str(i)
                style_cell_paragraph(cells[0], align="center", font="Times New Roman", size=10)
            if len(cells) >= 2:
                cells[1].text = (it.get("jenis_limbah") or "").strip()
                style_cell_paragraph(cells[1], align="left", left_indent_pt=6, font="Times New Roman", size=10)
            if len(cells) >= 3:
                cells[2].text = (it.get("kode_limbah") or "").strip()
                style_cell_paragraph(cells[2], align="center", font="Times New Roman", size=10)

    try:
        folder = str(FILES_DIR)
    except Exception:
        folder = "static/files"
    os.makedirs(folder, exist_ok=True)
    out_path = os.path.join(folder, f"{fname_base}.docx")
    doc.save(out_path)
    return f"{fname_base}.docx"
