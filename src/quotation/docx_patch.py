import os
import re
from datetime import datetime

from docx import Document  # ✅ NEW: untuk patch tahun di hasil docx

# ✅ OPTIONAL: kalau project kamu punya FILES_DIR (seperti invoice/mou kamu), pakai
try:
    from config_new import FILES_DIR
except Exception:
    FILES_DIR = None


# =========================
# ✅ NEW: Patch Tahun di hasil DOCX (tanggal + nomor surat)
# =========================

def _resolve_docx_path(filename: str) -> str:
    """
    Cari lokasi file docx dari filename yang dikembalikan create_docx().
    Coba:
    - kalau sudah absolute / ada di path sekarang
    - FILES_DIR (kalau ada)
    - static/files
    """
    if not filename:
        return ""

    # sudah path penuh / relative valid
    if os.path.exists(filename):
        return filename

    # coba FILES_DIR
    if FILES_DIR:
        try:
            p = os.path.join(str(FILES_DIR), filename)
            if os.path.exists(p):
                return p
        except Exception:
            pass

    # fallback default project kamu sering pakai static/files
    p2 = os.path.join("static", "files", filename)
    if os.path.exists(p2):
        return p2

    # terakhir: coba current dir join
    p3 = os.path.join(os.getcwd(), filename)
    if os.path.exists(p3):
        return p3

    return ""


def patch_docx_year_only(filename: str) -> None:
    """
    Memaksa tahun pada template DOCX ikut tahun sekarang.
    - Ganti tahun pada baris tanggal: "Tangerang, 2 Januari 2025" -> "... 2026"
    - Ganti tahun pada nomor surat: ".../IX/2025" -> ".../IX/2026"
    - Fallback: jika masih ada "2025" hardcode, ganti jadi tahun sekarang (aman).
    """
    full_path = _resolve_docx_path(filename)
    if not full_path:
        return

    tahun_now = str(datetime.now().year)

    # cocok untuk tanggal: "Kota, 2 Januari 2025"
    date_line_regex = re.compile(
        r"(\b[A-Za-zÀ-ÿ\.\-\s]+,\s*\d{1,2}\s+"
        r"(Januari|Februari|Maret|April|Mei|Juni|Juli|Agustus|September|Oktober|November|Desember)"
        r"\s+)(\d{4})\b",
        flags=re.IGNORECASE
    )

    # cocok untuk nomor surat seperti: "/IX/2025" atau "/XII/2025"
    nomor_regex = re.compile(r"(/(?:I|II|III|IV|V|VI|VII|VIII|IX|X|XI|XII)/)(\d{4})\b")

    def _fix_text(t: str) -> str:
        if not t:
            return t

        # 1) perbaiki tahun di baris tanggal
        t2 = date_line_regex.sub(lambda m: m.group(1) + tahun_now, t)

        # 2) perbaiki tahun di nomor surat
        t3 = nomor_regex.sub(lambda m: m.group(1) + tahun_now, t2)

        # 3) fallback (kalau template hardcode 2025 masih tersisa)
        t4 = re.sub(r"\b2025\b", tahun_now, t3)

        return t4

    def _apply_to_paragraph(p):
        if not p or not p.text:
            return
        new_text = _fix_text(p.text)
        if new_text == p.text:
            return
        # aman walau run split
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(new_text)

    doc = Document(full_path)

    for p in doc.paragraphs:
        _apply_to_paragraph(p)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _apply_to_paragraph(p)

    doc.save(full_path)
