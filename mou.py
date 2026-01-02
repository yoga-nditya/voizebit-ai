import os
import json
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from limbah_database import (
    find_limbah_by_kode,
    find_limbah_by_jenis,
)
from utils import (
    db_insert_history, db_append_message, db_update_state,
    search_company_address, search_company_address_ai,
    create_pdf,
)
from config_new import FILES_DIR


# =========================
# Helpers (as-is)
# =========================

def is_non_b3_input(text: str) -> bool:
    if not text:
        return False
    t = text.strip().lower()
    norm = re.sub(r'[\s\-_]+', '', t)
    return norm in ("nonb3", "nonbii3") or norm.startswith("nonb3")

def make_unique_filename_base(base_name: str) -> str:
    base_name = (base_name or "").strip()
    if not base_name:
        base_name = "Dokumen"
    try:
        folder = str(FILES_DIR)
    except Exception:
        folder = "static/files"

    def exists_any(name: str) -> bool:
        return (
            os.path.exists(os.path.join(folder, f"{name}.docx")) or
            os.path.exists(os.path.join(folder, f"{name}.pdf")) or
            os.path.exists(os.path.join(folder, f"{name}.xlsx")) or
            os.path.exists(os.path.join(folder, name))
        )

    if not exists_any(base_name):
        return base_name

    i = 2
    while True:
        candidate = f"{base_name} ({i})"
        if not exists_any(candidate):
            return candidate
        i += 1

def month_to_roman(m: int) -> str:
    rom = {
        1: "I", 2: "II", 3: "III", 4: "IV",
        5: "V", 6: "VI", 7: "VII", 8: "VIII",
        9: "IX", 10: "X", 11: "XI", 12: "XII"
    }
    return rom.get(m, "")

def company_to_code(name: str) -> str:
    if not name:
        return "XXX"
    t = re.sub(r'[^A-Za-z0-9 ]+', ' ', name).strip()
    t = re.sub(r'\s+', ' ', t)
    parts = [p for p in t.split() if p.lower() not in ("pt", "pt.", "persero", "tbk")]
    if not parts:
        return "XXX"
    if len(parts) == 1:
        return (parts[0][:3]).upper().ljust(3, "X")
    code = "".join([p[0] for p in parts[:3]]).upper()
    return code.ljust(3, "X")

def build_mou_nomor_surat(mou_data: dict) -> str:
    no_depan = (mou_data.get("nomor_depan") or "").strip()
    kode_p1 = company_to_code((mou_data.get("pihak_pertama") or "").strip())
    kode_p2 = (mou_data.get("pihak_kedua_kode") or "STBJ").strip().upper()
    kode_p3 = (mou_data.get("pihak_ketiga_kode") or "").strip().upper()
    now = datetime.now()
    romawi = month_to_roman(now.month)
    tahun = str(now.year)
    if not kode_p3:
        kode_p3 = "XXX"
    return f"{no_depan}/PKPLNB3/{kode_p1}-{kode_p2}-{kode_p3}/{romawi}/{tahun}"

def format_tanggal_indonesia(dt: datetime) -> str:
    hari_map = {0: "Senin", 1: "Selasa", 2: "Rabu", 3: "Kamis", 4: "Jumat", 5: "Sabtu", 6: "Minggu"}
    bulan_map = {1: "Januari", 2: "Februari", 3: "Maret", 4: "April", 5: "Mei", 6: "Juni",
                 7: "Juli", 8: "Agustus", 9: "September", 10: "Oktober", 11: "November", 12: "Desember"}
    hari = hari_map.get(dt.weekday(), "")
    bulan = bulan_map.get(dt.month, "")
    return f"{hari}, tanggal {dt.day} {bulan} {dt.year}"

def set_run_font(run, font_name="Times New Roman", size=10, bold=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:cs'), font_name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


# =========================
# Address FIX: kalau hasil search berisi kalimat panjang/penjelasan -> "Di tempat"
# =========================
def _sanitize_company_address(addr: str) -> str:
    a = (addr or "").strip()
    if not a:
        return "Di tempat"

    low = a.lower()

    bad_patterns = [
        r"tidak\s*dapat\s+menentukan",
        r"tidak\s*bisa\s+menentukan",
        r"tidak\s*dapat\s+menemukan",
        r"tidak\s*bisa\s+menemukan",
        r"tidak\s*ditemukan",
        r"tidak\s*ketemu",
        r"tidak\s+ada\s+informasi",
        r"tidak\s+memiliki\s+informasi",
        r"saya\s+tidak\s+memiliki",
        r"informasi\s+yang\s+cukup",
        r"tidak\s+cukup\s+informasi",
        r"untuk\s+menentukan",
        r"nama\s+tersebut\s+terlalu\s+umum",
        r"terlalu\s+umum",
        r"tidak\s+spesifik",
        r"placeholder",
        r"nama\s+contoh",
        r"banyak\s+perusahaan.*nama\s+serupa",
        r"mungkin\s+menggunakan\s+nama\s+serupa",
        r"maaf",
        r"gagal",
        r"cannot\s+find",
        r"not\s+found",
        r"no\s+information",
        r"no\s+result",
        r"unknown",
    ]

    for p in bad_patterns:
        if re.search(p, low):
            return "Di tempat"

    # Jika paragraf panjang dan tidak terlihat seperti alamat, fallback
    if len(a) > 120 and not re.search(r"\b(jl|jalan|rt|rw|kec|kel|kab|kota|no\.?|blok|desa)\b", low):
        return "Di tempat"

    if len(a) > 250:
        return "Di tempat"

    return a


def resolve_company_address(company_name: str) -> str:
    addr = ""
    try:
        addr = (search_company_address(company_name) or "").strip()
    except:
        addr = ""
    addr = _sanitize_company_address(addr)
    if addr != "Di tempat":
        return addr

    try:
        addr2 = (search_company_address_ai(company_name) or "").strip()
    except:
        addr2 = ""
    return _sanitize_company_address(addr2)


# =========================
# FIX 1: replace yang tahan run ter-split
# =========================
def replace_in_runs_keep_format(paragraph, old: str, new: str):
    if not old or not paragraph.text:
        return False
    if old not in paragraph.text:
        return False

    changed = False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            changed = True
    if changed:
        return True

    full = paragraph.text
    replaced = full.replace(old, new)
    if replaced == full:
        return False

    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(replaced)
    return True

def replace_in_cell_keep_format(cell, old: str, new: str):
    changed = False
    for p in cell.paragraphs:
        if replace_in_runs_keep_format(p, old, new):
            changed = True
    return changed

def replace_everywhere_keep_format(doc, old_list, new_value):
    if not new_value:
        return
    for p in doc.paragraphs:
        for old in old_list:
            replace_in_runs_keep_format(p, old, new_value)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for old in old_list:
                    replace_in_cell_keep_format(cell, old, new_value)

def style_cell_paragraph(cell, align="left", left_indent_pt=0, font="Times New Roman", size=10):
    if not cell.paragraphs:
        cell.add_paragraph("")
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if left_indent_pt and align == "left":
        p.paragraph_format.left_indent = Pt(left_indent_pt)
    for r in p.runs:
        set_run_font(r, font, size)

# =========================
# NEW: helper untuk center PIHAK KETIGA (kolom kanan)
# =========================
def _center_paragraph_if_contains(paragraph, needles):
    if not paragraph or not paragraph.text:
        return False
    txt = paragraph.text.strip()
    if not txt:
        return False
    for n in needles:
        if n and (n in txt):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            try:
                paragraph.paragraph_format.tab_stops.clear_all()
            except Exception:
                pass
            try:
                for r in paragraph.runs:
                    if r.text and r.text.startswith("\t"):
                        r.text = r.text.lstrip("\t")
            except Exception:
                pass
            return True
    return False

def _center_everywhere_for_needles(doc, needles):
    for p in doc.paragraphs:
        _center_paragraph_if_contains(p, needles)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _center_paragraph_if_contains(p, needles)


# =========================
# MoU Counter (as-is)
# =========================

def _mou_counter_path() -> str:
    try:
        folder = str(FILES_DIR)
    except Exception:
        folder = "static/files"
    os.makedirs(folder, exist_ok=True)
    return os.path.join(folder, "mou_counter.json")

def load_mou_counter() -> int:
    path = _mou_counter_path()
    try:
        if not os.path.exists(path):
            return -1
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f) or {}
        return int(data.get("counter", -1))
    except:
        return -1

def save_mou_counter(n: int) -> None:
    path = _mou_counter_path()
    with open(path, "w", encoding="utf-8") as f:
        json.dump({"counter": int(n)}, f)

def get_next_mou_no_depan() -> str:
    n = load_mou_counter() + 1
    save_mou_counter(n)
    return str(n).zfill(3)


# =========================
# Create MoU DOCX (as-is)
# =========================

def create_mou_docx(mou_data: dict, fname_base: str) -> str:
    template_path = "tamplate MoU.docx"
    if not os.path.exists(template_path):
        raise Exception("Template MoU tidak ditemukan. Pastikan file 'tamplate MoU.docx' ada di root project.")

    doc = Document(template_path)

    pihak1 = (mou_data.get("pihak_pertama") or "").strip()
    pihak2 = (mou_data.get("pihak_kedua") or "").strip()
    pihak3 = (mou_data.get("pihak_ketiga") or "").strip()

    alamat1 = (mou_data.get("alamat_pihak_pertama") or "").strip()
    alamat3 = (mou_data.get("alamat_pihak_ketiga") or "").strip()

    ttd1 = (mou_data.get("ttd_pihak_pertama") or "").strip()
    jab1 = (mou_data.get("jabatan_pihak_pertama") or "").strip()
    ttd3 = (mou_data.get("ttd_pihak_ketiga") or "").strip()
    jab3 = (mou_data.get("jabatan_pihak_ketiga") or "").strip()

    nomor_full = (mou_data.get("nomor_surat") or "").strip()
    tanggal_text = format_tanggal_indonesia(datetime.now())

    contoh_pihak1_candidates = [
        "PT. PANPAN LUCKY INDONESIA",
        "PT. Panpan Lucky Indonesia",
        "PT PANPAN LUCKY INDONESIA",
        "PT Panpan Lucky Indonesia",
    ]
    contoh_pihak2_candidates = [
        "PT. SARANA TRANS BERSAMA JAYA",
        "PT Sarana Trans Bersama Jaya",
        "PT SARANA TRANS BERSAMA JAYA",
        "PT Sarana Trans Bersama Jaya",
    ]
    contoh_pihak3_candidates = [
        "PT. HARAPAN BARU SEJAHTERA PLASTIK",
        "PT Harapan Baru Sejahtera Plastik",
        "PT HARAPAN BARU SEJAHTERA PLASTIK",
        "PT Harapan Baru Sejahtera Plastik",
    ]

    replace_everywhere_keep_format(doc, contoh_pihak1_candidates, pihak1.upper() if pihak1 else "")
    replace_everywhere_keep_format(doc, contoh_pihak2_candidates, pihak2.upper() if pihak2 else "")
    replace_everywhere_keep_format(doc, contoh_pihak3_candidates, pihak3.upper() if pihak3 else "")

    def replace_no_line(container_paragraphs):
        for p in container_paragraphs:
            if re.search(r'\bNo\s*:', p.text, flags=re.IGNORECASE):
                for run in p.runs:
                    if re.search(r'\bNo\s*:', run.text, flags=re.IGNORECASE):
                        run.text = re.sub(r'\bNo\s*:\s*.*', f"No : {nomor_full}", run.text, flags=re.IGNORECASE)
                        return True
                if p.runs:
                    p.runs[0].text = f"No : {nomor_full}"
                    for r in p.runs[1:]:
                        r.text = ""
                    return True
        return False

    replace_no_line(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if replace_no_line(cell.paragraphs):
                    break

    kalimat_tanggal = f"Pada hari ini {tanggal_text} kami yang bertanda tangan di bawah ini :"
    def replace_pada_hari_ini(container_paragraphs):
        for p in container_paragraphs:
            if "Pada hari ini" in p.text and "bertanda tangan" in p.text:
                if p.runs:
                    p.runs[0].text = kalimat_tanggal
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.add_run(kalimat_tanggal)
                return True
        return False

    replace_pada_hari_ini(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if replace_pada_hari_ini(cell.paragraphs):
                    break

    if ttd1:
        replace_everywhere_keep_format(doc, ["Huang Feifang"], ttd1)
    if jab1:
        replace_everywhere_keep_format(doc, ["Direktur Utama"], jab1)

    contoh_alamat_p1_candidates = [
        "Jl. Raya Serang KM. 22 No. 30, Desa Pasir Bolang, Kec Tigaraksa, Tangerang Banten",
        "Jl. Raya Serang KM. 22 No. 30, Desa Pasir Bolang, Kec. Tigaraksa, Tangerang Banten",
    ]
    if alamat1:
        replace_everywhere_keep_format(doc, contoh_alamat_p1_candidates, alamat1)

    if ttd3:
        replace_everywhere_keep_format(doc, ["Yogi Aditya", "Yogi Permana", "Yogi"], ttd3)
    if jab3:
        replace_everywhere_keep_format(doc, ["General Manager", "GENERAL MANAGER"], jab3)
        replace_everywhere_keep_format(doc, ["Direktur", "DIREKTUR"], jab3)

    contoh_alamat_p3_candidates = [
        "Jl. Karawang – Bekasi KM. 1 Bojongsari, Kec. Kedungwaringin, Kab. Bekasi – Jawa Barat",
        "Jl. Karawang - Bekasi KM. 1 Bojongsari, Kec. Kedungwaringin, Kab. Bekasi - Jawa Barat",
    ]
    if alamat3:
        replace_everywhere_keep_format(doc, contoh_alamat_p3_candidates, alamat3)

    pihak_ketiga_needles = [
        "Yogi Aditya", "Yogi Permana", "Yogi",
        "General Manager", "GENERAL MANAGER",
        "Direktur", "DIREKTUR",
        ttd3, jab3
    ]
    _center_everywhere_for_needles(doc, pihak_ketiga_needles)

    items = mou_data.get("items_limbah") or []
    target_table = None
    for t in doc.tables:
        if not t.rows:
            continue
        header_text = " ".join([c.text.strip() for c in t.rows[0].cells])
        if ("Jenis Limbah" in header_text) and ("Kode Limbah" in header_text):
            target_table = t
            break

    if target_table is not None:
        while len(target_table.rows) > 1:
            target_table._tbl.remove(target_table.rows[1]._tr)
        for i, it in enumerate(items, start=1):
            row = target_table.add_row()
            cells = row.cells
            if len(cells) >= 1:
                cells[0].text = str(i)
                style_cell_paragraph(cells[0], align="center", font="Times New Roman", size=10)
            if len(cells) >= 2:
                cells[1].text = (it.get("jenis_limbah") or "").strip()
                style_cell_paragraph(cells[1], align="left", left_indent_pt=6, font="Times New Roman", size=10)
            if len(cells) >= 3:
                cells[2].text = (it.get("kode_limbah") or "").strip()
                style_cell_paragraph(cells[2], align="center", font="Times New Roman", size=10)

    try:
        folder = str(FILES_DIR)
    except Exception:
        folder = "static/files"
    os.makedirs(folder, exist_ok=True)
    out_path = os.path.join(folder, f"{fname_base}.docx")
    doc.save(out_path)
    return f"{fname_base}.docx"


# =========================
# CHAT HANDLER MoU
# =========================

def handle_mou_flow(data: dict, text: str, lower: str, sid: str, state: dict, conversations: dict, history_id_in):
    """
    Return:
      - None jika bukan flow mou
      - dict response jika handled
    """

    if ('mou' in lower) and (state.get('step') == 'idle'):
        nomor_depan = get_next_mou_no_depan()
        state['step'] = 'mou_pihak_pertama'
        state['data'] = {
            'nomor_depan': nomor_depan,
            'nomor_surat': "",
            'items_limbah': [],
            'current_item': {},
            'pihak_kedua': "PT Sarana Trans Bersama Jaya",
            'pihak_kedua_kode': "STBJ",
            'pihak_pertama': "",
            'alamat_pihak_pertama': "",
            'pihak_ketiga': "",
            'pihak_ketiga_kode': "",
            'alamat_pihak_ketiga': "",
            'ttd_pihak_pertama': "",
            'jabatan_pihak_pertama': "",
            'ttd_pihak_ketiga': "",
            'jabatan_pihak_ketiga': "",
        }
        conversations[sid] = state

        out_text = (
            "Baik, saya akan membantu menyusun <b>MoU</b>.<br><br>"
        )

        history_id_created = None
        if not history_id_in:
            history_id_created = db_insert_history(
                title="Chat Baru",
                task_type=data.get("taskType") or "mou",
                data={},
                files=[],
                messages=[
                    {"id": __import__("uuid").uuid4().hex[:12], "sender": "user", "text": text, "files": [], "timestamp": datetime.now().isoformat()},
                    {"id": __import__("uuid").uuid4().hex[:12], "sender": "assistant", "text": re.sub(r'<br\s*/?>', '\n', out_text), "files": [], "timestamp": datetime.now().isoformat()},
                ],
                state=state
            )
        else:
            db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
            db_update_state(int(history_id_in), state)

        return {"text": out_text, "history_id": history_id_created or history_id_in}

    if state.get('step') == 'mou_pihak_pertama':
        state['data']['pihak_pertama'] = text.strip()

        # ✅ FIX: pakai resolver yang otomatis "Di tempat" kalau hasil panjang/penjelasan
        alamat = resolve_company_address(text)
        state['data']['alamat_pihak_pertama'] = alamat

        state['step'] = 'mou_pilih_pihak_ketiga'
        conversations[sid] = state

        out_text = (
            f"Pihak Pertama: <b>{state['data']['pihak_pertama']}</b><br>"
            f"Alamat: <b>{alamat}</b><br><br>"
            "Pertanyaan 2: <b>Pilih Pihak Ketiga (Pengelola Limbah)</b><br>"
            "1. HBSP<br>2. KJL<br>3. MBI<br>4. CGA<br><br>"
            "<i>Ketik nomor 1-4 atau ketik HBSP/KJL/MBI/CGA</i>"
        )

        if history_id_in:
            db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
            db_update_state(int(history_id_in), state)

        return {"text": out_text, "history_id": history_id_in}

    if state.get('step') == 'mou_pilih_pihak_ketiga':
        pilihan = text.strip().upper()
        mapping = {"1": "HBSP", "2": "KJL", "3": "MBI", "4": "CGA", "HBSP": "HBSP", "KJL": "KJL", "MBI": "MBI", "CGA": "CGA"}
        kode = mapping.get(pilihan)
        if not kode:
            out_text = (
                "Input tidak valid.<br><br>"
                "Pilih Pihak Ketiga:<br>"
                "1. HBSP<br>2. KJL<br>3. MBI<br>4. CGA"
            )
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
            return {"text": out_text, "history_id": history_id_in}

        pihak3_nama_map = {"HBSP": "PT Harapan Baru Sejahtera Plastik", "KJL": "KJL", "MBI": "MBI", "CGA": "CGA"}
        pihak3_alamat_map = {
            "HBSP": "Jl. Karawang – Bekasi KM. 1 Bojongsari, Kec. Kedungwaringin, Kab. Bekasi – Jawa Barat",
            "KJL": "",
            "MBI": "",
            "CGA": "",
        }

        state['data']['pihak_ketiga'] = pihak3_nama_map.get(kode, kode)
        state['data']['pihak_ketiga_kode'] = kode
        state['data']['alamat_pihak_ketiga'] = pihak3_alamat_map.get(kode, "")

        state['data']['nomor_surat'] = build_mou_nomor_surat(state['data'])
        state['step'] = 'mou_jenis_kode_limbah'
        state['data']['current_item'] = {}
        conversations[sid] = state

        out_text = (
            f"Pihak Ketiga: <b>{state['data']['pihak_ketiga']}</b><br>"
            f"Nomor MoU: <b>{state['data']['nomor_surat']}</b><br><br>"
            "Item 1<br>"
            "Pertanyaan 3: <b>Sebutkan jenis limbah atau kode limbah</b><br>"
            "<i>Contoh: A102d atau aki baterai bekas. Atau ketik NON B3 untuk input manual.</i>"
        )

        if history_id_in:
            db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
            db_update_state(int(history_id_in), state)

        return {"text": out_text, "history_id": history_id_in}

    if state.get('step') == 'mou_jenis_kode_limbah':
        if is_non_b3_input(text):
            state['data']['current_item']['kode_limbah'] = "NON B3"
            state['data']['current_item']['jenis_limbah'] = ""
            state['step'] = 'mou_manual_jenis_limbah'
            conversations[sid] = state
            out_text = (
                "Kode limbah: <b>NON B3</b><br><br>"
                "Pertanyaan 3A: <b>Jenis limbah (manual) apa?</b>"
            )
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return {"text": out_text, "history_id": history_id_in}

        kode, data_limbah = find_limbah_by_kode(text)
        if not (kode and data_limbah):
            kode, data_limbah = find_limbah_by_jenis(text)

        if kode and data_limbah:
            state['data']['current_item']['kode_limbah'] = kode
            state['data']['current_item']['jenis_limbah'] = data_limbah['jenis']
            state['data']['items_limbah'].append(state['data']['current_item'])
            num = len(state['data']['items_limbah'])
            state['step'] = 'mou_tambah_item'
            state['data']['current_item'] = {}
            conversations[sid] = state
            out_text = (
                f"Item {num} tersimpan.<br>"
                f"Jenis: <b>{data_limbah['jenis']}</b><br>"
                f"Kode: <b>{kode}</b><br><br>"
                "Pertanyaan: <b>Tambah item lagi?</b> (ya/tidak)"
            )
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return {"text": out_text, "history_id": history_id_in}

        out_text = (
            f"Maaf, limbah <b>{text}</b> tidak ditemukan.<br><br>"
            "Silakan ketik kode/jenis lain atau ketik <b>NON B3</b> untuk input manual."
        )
        if history_id_in:
            db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
            db_update_state(int(history_id_in), state)
        return {"text": out_text, "history_id": history_id_in}

    if state.get('step') == 'mou_manual_jenis_limbah':
        state['data']['current_item']['jenis_limbah'] = text.strip()
        state['data']['items_limbah'].append(state['data']['current_item'])
        num = len(state['data']['items_limbah'])
        state['step'] = 'mou_tambah_item'
        state['data']['current_item'] = {}
        conversations[sid] = state
        out_text = (
            f"Item {num} tersimpan.<br>"
            f"Jenis (manual): <b>{state['data']['items_limbah'][-1]['jenis_limbah']}</b><br>"
            "Kode: <b>NON B3</b><br><br>"
            "Pertanyaan: <b>Tambah item lagi?</b> (ya/tidak)"
        )
        if history_id_in:
            db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
            db_update_state(int(history_id_in), state)
        return {"text": out_text, "history_id": history_id_in}

    if state.get('step') == 'mou_tambah_item':
        if re.match(r'^\d+', text.strip()):
            out_text = (
                "Mohon jawab dengan <b>ya</b> atau <b>tidak</b>.<br><br>"
                "Pertanyaan: <b>Tambah item lagi?</b>"
            )
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
            return {"text": out_text, "history_id": history_id_in}

        if ('ya' in lower) or ('iya' in lower):
            num = len(state['data']['items_limbah'])
            state['step'] = 'mou_jenis_kode_limbah'
            state['data']['current_item'] = {}
            conversations[sid] = state
            out_text = (
                f"Item {num+1}<br>"
                "Pertanyaan 3: <b>Sebutkan jenis limbah atau kode limbah</b><br>"
                "<i>Contoh: A102d. Atau ketik NON B3 untuk input manual.</i>"
            )
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return {"text": out_text, "history_id": history_id_in}

        if ('tidak' in lower) or ('skip' in lower) or ('lewat' in lower) or ('gak' in lower) or ('nggak' in lower):
            state['step'] = 'mou_ttd_pihak_pertama'
            conversations[sid] = state
            out_text = "Pertanyaan: <b>Nama penandatangan Pihak Pertama?</b>"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return {"text": out_text, "history_id": history_id_in}

        out_text = (
            "Mohon jawab dengan <b>ya</b> atau <b>tidak</b>.<br><br>"
            "Pertanyaan: <b>Tambah item lagi?</b>"
        )
        if history_id_in:
            db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
        return {"text": out_text, "history_id": history_id_in}

    if state.get('step') == 'mou_ttd_pihak_pertama':
        state['data']['ttd_pihak_pertama'] = text.strip()
        state['step'] = 'mou_jabatan_pihak_pertama'
        conversations[sid] = state
        out_text = "Pertanyaan: <b>Jabatan penandatangan Pihak Pertama?</b>"
        if history_id_in:
            db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
            db_update_state(int(history_id_in), state)
        return {"text": out_text, "history_id": history_id_in}

    if state.get('step') == 'mou_jabatan_pihak_pertama':
        state['data']['jabatan_pihak_pertama'] = text.strip()
        state['step'] = 'mou_ttd_pihak_ketiga'
        conversations[sid] = state
        out_text = "Pertanyaan: <b>Nama penandatangan Pihak Ketiga?</b>"
        if history_id_in:
            db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
            db_update_state(int(history_id_in), state)
        return {"text": out_text, "history_id": history_id_in}

    if state.get('step') == 'mou_ttd_pihak_ketiga':
        state['data']['ttd_pihak_ketiga'] = text.strip()
        state['step'] = 'mou_jabatan_pihak_ketiga'
        conversations[sid] = state
        out_text = "Pertanyaan: <b>Jabatan penandatangan Pihak Ketiga?</b>"
        if history_id_in:
            db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
            db_update_state(int(history_id_in), state)
        return {"text": out_text, "history_id": history_id_in}

    if state.get('step') == 'mou_jabatan_pihak_ketiga':
        state['data']['jabatan_pihak_ketiga'] = text.strip()

        nama_pt_raw = state['data'].get('pihak_pertama', '').strip()
        safe_pt = re.sub(r'[^A-Za-z0-9 \-]+', '', nama_pt_raw).strip()
        safe_pt = re.sub(r'\s+', ' ', safe_pt).strip()

        base_fname = f"MoU - {safe_pt}" if safe_pt else "MoU - Perusahaan"
        fname_base = make_unique_filename_base(base_fname)

        if not state['data'].get("nomor_surat"):
            state['data']['nomor_surat'] = build_mou_nomor_surat(state['data'])

        docx = create_mou_docx(state['data'], fname_base)
        pdf = create_pdf(fname_base)

        conversations[sid] = {'step': 'idle', 'data': {}}

        files = [{"type": "docx", "filename": docx, "url": f"/download/{docx}"}]
        if pdf:
            files.append({"type": "pdf", "filename": pdf, "url": f"/download/{pdf}"})

        history_title = f"MoU {nama_pt_raw}" if nama_pt_raw else "MoU"
        history_task_type = "mou"

        if history_id_in:
            from utils import db_connect
            conn = db_connect()
            cur = conn.cursor()
            cur.execute("""
                UPDATE chat_history
                SET title = ?, task_type = ?, data_json = ?, files_json = ?
                WHERE id = ?
            """, (
                history_title,
                history_task_type,
                json.dumps(state['data'], ensure_ascii=False),
                json.dumps(files, ensure_ascii=False),
                int(history_id_in),
            ))
            conn.commit()
            conn.close()
            history_id = int(history_id_in)
        else:
            history_id = db_insert_history(
                title=history_title,
                task_type=history_task_type,
                data=state['data'],
                files=files,
                messages=[],
                state={}
            )

        out_text = (
            "<b>MoU berhasil dibuat.</b><br><br>"
        )

        db_append_message(history_id, "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=files)
        return {"text": out_text, "files": files, "history_id": history_id}

    return None
