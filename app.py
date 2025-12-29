import os
import json
import uuid
import re
import platform
from datetime import datetime

from flask import Flask, request, jsonify, render_template, send_from_directory, session

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

from config_new import *
from limbah_database import (
    LIMBAH_B3_DB,
    find_limbah_by_kode,
    find_limbah_by_jenis,
    convert_voice_to_number,
    parse_termin_days,
    angka_ke_terbilang,
    format_rupiah
)
from utils import (
    init_db, load_counter,
    db_insert_history, db_list_histories, db_get_history_detail,
    db_update_title, db_delete_history, db_append_message, db_update_state,
    get_next_nomor, create_docx, create_pdf,
    search_company_address, search_company_address_ai, call_ai,
    PDF_AVAILABLE, PDF_METHOD, LIBREOFFICE_PATH
)

app = Flask(__name__, static_folder="static", template_folder="templates")
app.secret_key = FLASK_SECRET_KEY

conversations = {}
init_db()

@app.after_request
def add_cors_headers(resp):
    resp.headers["Access-Control-Allow-Origin"] = "*"
    resp.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization"
    resp.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS"
    return resp

def is_non_b3_input(text: str) -> bool:
    if not text:
        return False
    t = text.strip().lower()
    norm = re.sub(r'[\s\-_]+', '', t)
    return norm in ("nonb3", "nonbii3") or norm.startswith("nonb3")

def normalize_id_number_text(text: str) -> str:
    if not text:
        return text
    t = text.strip()
    t = re.sub(r'(?<=\d)[\.,](?=\d{3}(\D|$))', '', t)
    t = re.sub(r'(?<=\d),(?=\d)', '.', t)
    return t

def parse_amount_id(text: str) -> int:
    if not text:
        return 0
    raw = text.strip()
    lower = raw.lower()

    digit_map = {
        "nol": 0, "kosong": 0,
        "satu": 1, "se": 1,
        "dua": 2,
        "tiga": 3,
        "empat": 4,
        "lima": 5,
        "enam": 6,
        "tujuh": 7,
        "delapan": 8,
        "sembilan": 9
    }

    def token_to_digit(tok: str):
        tok = tok.strip().lower()
        if tok.isdigit():
            return int(tok)
        return digit_map.get(tok, None)

    scale_map = {
        "ribu": 1_000,
        "juta": 1_000_000,
        "miliar": 1_000_000_000,
        "triliun": 1_000_000_000_000,
    }
    scale = None
    for k, m in scale_map.items():
        if re.search(rf'\b{k}\b', lower):
            scale = m
            break

    if "koma" in lower:
        parts = re.split(r'\bkoma\b', lower, maxsplit=1)
        left_part = parts[0].strip()
        right_part = parts[1].strip() if len(parts) > 1 else ""
        left_tokens = re.findall(r'[a-zA-Z0-9]+', left_part)
        right_tokens = re.findall(r'[a-zA-Z0-9]+', right_part)
        left_digit = token_to_digit(left_tokens[-1]) if left_tokens else None
        right_digit = token_to_digit(right_tokens[0]) if right_tokens else None
        if left_digit is not None and right_digit is not None:
            val = float(f"{left_digit}.{right_digit}")
            if scale:
                val *= scale
            return int(round(val))

    tnorm = normalize_id_number_text(raw)
    val = convert_voice_to_number(tnorm)
    if val is None:
        val = 0

    try:
        f = float(val)
        if scale and f < scale:
            val = f * scale
    except:
        pass

    try:
        return int(round(float(val)))
    except:
        digits = re.sub(r'\D+', '', str(val))
        return int(digits) if digits else 0

def parse_qty_id(text: str) -> float:
    if not text:
        return 0.0
    t = normalize_id_number_text(text)
    v = convert_voice_to_number(t)
    try:
        return float(v)
    except:
        m = re.findall(r'\d+(?:\.\d+)?', t)
        return float(m[0]) if m else 0.0

def make_unique_filename_base(base_name: str) -> str:
    base_name = (base_name or "").strip()
    if not base_name:
        base_name = "Dokumen"
    try:
        folder = str(FILES_DIR)
    except Exception:
        folder = "static/files"

    def exists_any(name: str) -> bool:
        return (
            os.path.exists(os.path.join(folder, f"{name}.docx")) or
            os.path.exists(os.path.join(folder, f"{name}.pdf")) or
            os.path.exists(os.path.join(folder, f"{name}.xlsx")) or
            os.path.exists(os.path.join(folder, name))
        )

    if not exists_any(base_name):
        return base_name

    i = 2
    while True:
        candidate = f"{base_name} ({i})"
        if not exists_any(candidate):
            return candidate
        i += 1

def _mou_counter_path() -> str:
    try:
        folder = str(FILES_DIR)
    except Exception:
        folder = "static/files"
    os.makedirs(folder, exist_ok=True)
    return os.path.join(folder, "mou_counter.json")

def load_mou_counter() -> int:
    path = _mou_counter_path()
    try:
        if not os.path.exists(path):
            return -1
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f) or {}
        return int(data.get("counter", -1))
    except:
        return -1

def save_mou_counter(n: int) -> None:
    path = _mou_counter_path()
    with open(path, "w", encoding="utf-8") as f:
        json.dump({"counter": int(n)}, f)

def get_next_mou_no_depan() -> str:
    n = load_mou_counter() + 1
    save_mou_counter(n)
    return str(n).zfill(3)

def _invoice_counter_path() -> str:
    try:
        folder = str(FILES_DIR)
    except Exception:
        folder = "static/files"
    os.makedirs(folder, exist_ok=True)
    return os.path.join(folder, "invoice_counter.json")

def load_invoice_counter() -> int:
    path = _invoice_counter_path()
    try:
        if not os.path.exists(path):
            return 0
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f) or {}
        return int(data.get("counter", 0))
    except:
        return 0

def save_invoice_counter(n: int) -> None:
    path = _invoice_counter_path()
    with open(path, "w", encoding="utf-8") as f:
        json.dump({"counter": int(n)}, f)

def get_next_invoice_no() -> str:
    now = datetime.now()
    prefix = now.strftime("%y%m")
    n = load_invoice_counter() + 1
    save_invoice_counter(n)
    return f"{prefix}{str(n).zfill(3)}"

def month_to_roman(m: int) -> str:
    rom = {
        1: "I", 2: "II", 3: "III", 4: "IV",
        5: "V", 6: "VI", 7: "VII", 8: "VIII",
        9: "IX", 10: "X", 11: "XI", 12: "XII"
    }
    return rom.get(m, "")

def company_to_code(name: str) -> str:
    if not name:
        return "XXX"
    t = re.sub(r'[^A-Za-z0-9 ]+', ' ', name).strip()
    t = re.sub(r'\s+', ' ', t)
    parts = [p for p in t.split() if p.lower() not in ("pt", "pt.", "persero", "tbk")]
    if not parts:
        return "XXX"
    if len(parts) == 1:
        return (parts[0][:3]).upper().ljust(3, "X")
    code = "".join([p[0] for p in parts[:3]]).upper()
    return code.ljust(3, "X")

def build_mou_nomor_surat(mou_data: dict) -> str:
    no_depan = (mou_data.get("nomor_depan") or "").strip()
    kode_p1 = company_to_code((mou_data.get("pihak_pertama") or "").strip())
    kode_p2 = (mou_data.get("pihak_kedua_kode") or "STBJ").strip().upper()
    kode_p3 = (mou_data.get("pihak_ketiga_kode") or "").strip().upper()
    now = datetime.now()
    romawi = month_to_roman(now.month)
    tahun = str(now.year)
    if not kode_p3:
        kode_p3 = "XXX"
    return f"{no_depan}/PKPLNB3/{kode_p1}-{kode_p2}-{kode_p3}/{romawi}/{tahun}"

def format_tanggal_indonesia(dt: datetime) -> str:
    hari_map = {0: "Senin", 1: "Selasa", 2: "Rabu", 3: "Kamis", 4: "Jumat", 5: "Sabtu", 6: "Minggu"}
    bulan_map = {1: "Januari", 2: "Februari", 3: "Maret", 4: "April", 5: "Mei", 6: "Juni",
                 7: "Juli", 8: "Agustus", 9: "September", 10: "Oktober", 11: "November", 12: "Desember"}
    hari = hari_map.get(dt.weekday(), "")
    bulan = bulan_map.get(dt.month, "")
    return f"{hari}, tanggal {dt.day} {bulan} {dt.year}"

def set_run_font(run, font_name="Times New Roman", size=10, bold=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:cs'), font_name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold

def replace_in_runs_keep_format(paragraph, old: str, new: str):
    if not old or not paragraph.text:
        return False
    if old not in paragraph.text:
        return False
    changed = False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            changed = True
    return changed

def replace_in_cell_keep_format(cell, old: str, new: str):
    changed = False
    for p in cell.paragraphs:
        if replace_in_runs_keep_format(p, old, new):
            changed = True
    return changed

def replace_everywhere_keep_format(doc, old_list, new_value):
    if not new_value:
        return
    for p in doc.paragraphs:
        for old in old_list:
            replace_in_runs_keep_format(p, old, new_value)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for old in old_list:
                    replace_in_cell_keep_format(cell, old, new_value)

def style_cell_paragraph(cell, align="left", left_indent_pt=0, font="Times New Roman", size=10):
    if not cell.paragraphs:
        cell.add_paragraph("")
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if left_indent_pt and align == "left":
        p.paragraph_format.left_indent = Pt(left_indent_pt)
    for r in p.runs:
        set_run_font(r, font, size)

def create_mou_docx(mou_data: dict, fname_base: str) -> str:
    template_path = "tamplate MoU.docx"
    if not os.path.exists(template_path):
        raise Exception("Template MoU tidak ditemukan. Pastikan file 'tamplate MoU.docx' ada di root project.")

    doc = Document(template_path)

    pihak1 = (mou_data.get("pihak_pertama") or "").strip()
    pihak2 = (mou_data.get("pihak_kedua") or "").strip()
    pihak3 = (mou_data.get("pihak_ketiga") or "").strip()

    alamat1 = (mou_data.get("alamat_pihak_pertama") or "").strip()
    alamat3 = (mou_data.get("alamat_pihak_ketiga") or "").strip()

    ttd1 = (mou_data.get("ttd_pihak_pertama") or "").strip()
    jab1 = (mou_data.get("jabatan_pihak_pertama") or "").strip()
    ttd3 = (mou_data.get("ttd_pihak_ketiga") or "").strip()
    jab3 = (mou_data.get("jabatan_pihak_ketiga") or "").strip()

    nomor_full = (mou_data.get("nomor_surat") or "").strip()
    tanggal_text = format_tanggal_indonesia(datetime.now())

    contoh_pihak1_candidates = [
        "PT. PANPAN LUCKY INDONESIA",
        "PT. Panpan Lucky Indonesia",
        "PT PANPAN LUCKY INDONESIA",
        "PT Panpan Lucky Indonesia",
    ]
    contoh_pihak2_candidates = [
        "PT. SARANA TRANS BERSAMA JAYA",
        "PT Sarana Trans Bersama Jaya",
        "PT SARANA TRANS BERSAMA JAYA",
        "PT Sarana Trans Bersama Jaya",
    ]
    contoh_pihak3_candidates = [
        "PT. HARAPAN BARU SEJAHTERA PLASTIK",
        "PT Harapan Baru Sejahtera Plastik",
        "PT HARAPAN BARU SEJAHTERA PLASTIK",
        "PT Harapan Baru Sejahtera Plastik",
    ]

    replace_everywhere_keep_format(doc, contoh_pihak1_candidates, pihak1.upper() if pihak1 else "")
    replace_everywhere_keep_format(doc, contoh_pihak2_candidates, pihak2.upper() if pihak2 else "")
    replace_everywhere_keep_format(doc, contoh_pihak3_candidates, pihak3.upper() if pihak3 else "")

    def replace_no_line(container_paragraphs):
        for p in container_paragraphs:
            if re.search(r'\bNo\s*:', p.text, flags=re.IGNORECASE):
                for run in p.runs:
                    if re.search(r'\bNo\s*:', run.text, flags=re.IGNORECASE):
                        run.text = re.sub(r'\bNo\s*:\s*.*', f"No : {nomor_full}", run.text, flags=re.IGNORECASE)
                        return True
                if p.runs:
                    p.runs[0].text = f"No : {nomor_full}"
                    for r in p.runs[1:]:
                        r.text = ""
                    return True
        return False

    replace_no_line(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if replace_no_line(cell.paragraphs):
                    break

    kalimat_tanggal = f"Pada hari ini {tanggal_text} kami yang bertanda tangan di bawah ini :"
    def replace_pada_hari_ini(container_paragraphs):
        for p in container_paragraphs:
            if "Pada hari ini" in p.text and "bertanda tangan" in p.text:
                if p.runs:
                    p.runs[0].text = kalimat_tanggal
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.add_run(kalimat_tanggal)
                return True
        return False

    replace_pada_hari_ini(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if replace_pada_hari_ini(cell.paragraphs):
                    break

    if ttd1:
        replace_everywhere_keep_format(doc, ["Huang Feifang"], ttd1)
    if jab1:
        replace_everywhere_keep_format(doc, ["Direktur Utama"], jab1)

    contoh_alamat_p1_candidates = [
        "Jl. Raya Serang KM. 22 No. 30, Desa Pasir Bolang, Kec Tigaraksa, Tangerang Banten",
        "Jl. Raya Serang KM. 22 No. 30, Desa Pasir Bolang, Kec. Tigaraksa, Tangerang Banten",
    ]
    if alamat1:
        replace_everywhere_keep_format(doc, contoh_alamat_p1_candidates, alamat1)

    if ttd3:
        replace_everywhere_keep_format(doc, ["Yogi Aditya", "Yogi Permana", "Yogi"], ttd3)
    if jab3:
        replace_everywhere_keep_format(doc, ["General Manager", "GENERAL MANAGER"], jab3)
        replace_everywhere_keep_format(doc, ["Direktur", "DIREKTUR"], jab3)

    contoh_alamat_p3_candidates = [
        "Jl. Karawang – Bekasi KM. 1 Bojongsari, Kec. Kedungwaringin, Kab. Bekasi – Jawa Barat",
        "Jl. Karawang - Bekasi KM. 1 Bojongsari, Kec. Kedungwaringin, Kab. Bekasi - Jawa Barat",
    ]
    if alamat3:
        replace_everywhere_keep_format(doc, contoh_alamat_p3_candidates, alamat3)

    items = mou_data.get("items_limbah") or []
    target_table = None
    for t in doc.tables:
        if not t.rows:
            continue
        header_text = " ".join([c.text.strip() for c in t.rows[0].cells])
        if ("Jenis Limbah" in header_text) and ("Kode Limbah" in header_text):
            target_table = t
            break

    if target_table is not None:
        while len(target_table.rows) > 1:
            target_table._tbl.remove(target_table.rows[1]._tr)
        for i, it in enumerate(items, start=1):
            row = target_table.add_row()
            cells = row.cells
            if len(cells) >= 1:
                cells[0].text = str(i)
                style_cell_paragraph(cells[0], align="center", font="Times New Roman", size=10)
            if len(cells) >= 2:
                cells[1].text = (it.get("jenis_limbah") or "").strip()
                style_cell_paragraph(cells[1], align="left", left_indent_pt=6, font="Times New Roman", size=10)
            if len(cells) >= 3:
                cells[2].text = (it.get("kode_limbah") or "").strip()
                style_cell_paragraph(cells[2], align="center", font="Times New Roman", size=10)

    try:
        folder = str(FILES_DIR)
    except Exception:
        folder = "static/files"
    os.makedirs(folder, exist_ok=True)
    out_path = os.path.join(folder, f"{fname_base}.docx")
    doc.save(out_path)
    return f"{fname_base}.docx"

def _thin_border():
    side = Side(style="thin", color="000000")
    return Border(left=side, right=side, top=side, bottom=side)

def _set_border(ws, r1, c1, r2, c2, border):
    for r in range(r1, r2 + 1):
        for c in range(c1, c2 + 1):
            ws.cell(r, c).border = border

def create_invoice_xlsx(inv: dict, fname_base: str) -> str:
    wb = Workbook()
    ws = wb.active
    ws.title = "Invoice"

    ws.page_setup.orientation = ws.ORIENTATION_PORTRAIT
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_margins.left = 0.35
    ws.page_margins.right = 0.35
    ws.page_margins.top = 0.35
    ws.page_margins.bottom = 0.35

    col_widths = {"A": 8, "B": 6, "C": 12, "D": 45, "E": 14, "F": 16}
    for col, w in col_widths.items():
        ws.column_dimensions[col].width = w

    border = _thin_border()
    bold = Font(bold=True)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="top", wrap_text=True)
    right = Alignment(horizontal="right", vertical="center", wrap_text=True)

    def money(cell):
        cell.number_format = '#,##0'

    payment = inv.get("payment") or {}
    defaults = {
        "beneficiary": "PT. Sarana Trans Bersama Jaya",
        "bank_name": "BCA",
        "branch": "Cibadak - Sukabumi",
        "idr_acct": "35212 26666",
    }
    for k, v in defaults.items():
        if not payment.get(k):
            payment[k] = v

    invoice_no = inv.get("invoice_no") or get_next_invoice_no()
    inv_date = inv.get("invoice_date") or datetime.now().strftime("%d-%b-%y")

    bill_to = inv.get("bill_to") or {}
    ship_to = inv.get("ship_to") or {}
    attn = inv.get("attn") or "Accounting / Finance"
    phone = inv.get("phone") or ""
    fax = inv.get("fax") or ""

    sales_person = inv.get("sales_person") or "Syaeful Bakri"
    ref_no = inv.get("ref_no") or ""
    ship_via = inv.get("ship_via") or ""
    ship_date = inv.get("ship_date") or ""
    terms = inv.get("terms") or ""
    no_surat_jalan = inv.get("no_surat_jalan") or ""

    ws["A1"].value = "Bill To:"
    ws["A1"].font = bold
    ws.merge_cells("A1:C1")

    ws["D1"].value = "Ship To:"
    ws["D1"].font = bold
    ws.merge_cells("D1:F1")

    bill_lines = [(bill_to.get("name") or "").strip(), (bill_to.get("address") or "").strip(), (bill_to.get("address2") or "").strip()]
    ship_lines = [(ship_to.get("name") or "").strip(), (ship_to.get("address") or "").strip(), (ship_to.get("address2") or "").strip()]
    bill_text = "\n".join([x for x in bill_lines if x])
    ship_text = "\n".join([x for x in ship_lines if x])

    ws["A2"].value = bill_text
    ws.merge_cells("A2:C3")
    ws["A2"].alignment = left

    ws["D2"].value = ship_text
    ws.merge_cells("D2:F3")
    ws["D2"].alignment = left

    ws["A4"].value = "Phone:"
    ws["A4"].font = bold
    ws.merge_cells("A4:B4")
    ws["C4"].value = phone
    ws["C4"].alignment = left

    ws["D4"].value = "Fax:"
    ws["D4"].font = bold
    ws.merge_cells("D4:E4")
    ws["F4"].value = fax
    ws["F4"].alignment = left

    ws["A5"].value = "Attn :"
    ws["A5"].font = bold
    ws.merge_cells("A5:B5")
    ws["C5"].value = attn
    ws.merge_cells("C5:F5")
    ws["C5"].alignment = left

    ws["E6"].value = "Invoice"
    ws["E6"].font = bold
    ws["E6"].alignment = center
    ws["F6"].value = invoice_no
    ws["F6"].alignment = center

    ws["E7"].value = "Date"
    ws["E7"].font = bold
    ws["E7"].alignment = center
    ws["F7"].value = inv_date
    ws["F7"].alignment = center

    ws["E8"].value = "No. Surat Jalan"
    ws["E8"].font = bold
    ws["E8"].alignment = center
    ws["F8"].value = no_surat_jalan
    ws["F8"].alignment = center

    ws.merge_cells("A9:B9")
    ws["A9"].value = "Ref No."
    ws["A9"].font = bold
    ws["A9"].alignment = center

    ws.merge_cells("C9:D9")
    ws["C9"].value = "Sales Person"
    ws["C9"].font = bold
    ws["C9"].alignment = center

    ws["E9"].value = "Ship Via"
    ws["E9"].font = bold
    ws["E9"].alignment = center

    ws["F9"].value = "Ship Date"
    ws["F9"].font = bold
    ws["F9"].alignment = center

    ws.merge_cells("A10:B10")
    ws["A10"].value = ref_no
    ws["A10"].alignment = center

    ws.merge_cells("C10:D10")
    ws["C10"].value = sales_person
    ws["C10"].alignment = center

    ws["E10"].value = ship_via
    ws["E10"].alignment = center

    ws["F10"].value = ship_date
    ws["F10"].alignment = center

    ws.merge_cells("E11:F11")
    ws["E11"].value = "Terms"
    ws["E11"].font = bold
    ws["E11"].alignment = center

    ws.merge_cells("E12:F12")
    ws["E12"].value = terms
    ws["E12"].alignment = center

    _set_border(ws, 1, 1, 12, 6, border)

    ws["A14"].value = "Qty"
    ws["B14"].value = ""
    ws["C14"].value = "Date"
    ws["D14"].value = "Description"
    ws["E14"].value = "Price"
    ws["F14"].value = "Amount (IDR)"
    for c in "ABCDEF":
        ws[f"{c}14"].font = bold
        ws[f"{c}14"].alignment = center

    items = inv.get("items") or []
    r = 15
    subtotal = 0
    for it in items:
        qty = float(it.get("qty") or 0)
        unit = (it.get("unit") or "").strip()
        desc = (it.get("description") or "").strip()
        price = int(it.get("price") or 0)
        line_date = it.get("date") or inv_date
        amount = int(round(qty * price))
        subtotal += amount

        ws[f"A{r}"].value = qty if qty % 1 != 0 else int(qty)
        ws[f"A{r}"].alignment = center
        ws[f"B{r}"].value = unit
        ws[f"B{r}"].alignment = center
        ws[f"C{r}"].value = line_date
        ws[f"C{r}"].alignment = center
        ws[f"D{r}"].value = desc
        ws[f"D{r}"].alignment = left
        ws[f"E{r}"].value = price
        ws[f"E{r}"].alignment = right
        money(ws[f"E{r}"])
        ws[f"F{r}"].value = amount
        ws[f"F{r}"].alignment = right
        money(ws[f"F{r}"])
        r += 1

    min_last_row = 26
    if r < min_last_row:
        r = min_last_row

    _set_border(ws, 14, 1, r - 1, 6, border)

    freight = int(inv.get("freight") or 0)
    ppn_rate = float(inv.get("ppn_rate") or 0.11)
    deposit = int(inv.get("deposit") or 0)

    total_before_ppn = subtotal + freight
    ppn = int(round(total_before_ppn * ppn_rate))
    balance = total_before_ppn + ppn - deposit

    sum_row = r
    ws[f"E{sum_row}"].value = "Total"
    ws[f"E{sum_row}"].alignment = right
    ws[f"E{sum_row}"].font = bold
    ws[f"F{sum_row}"].value = subtotal
    ws[f"F{sum_row}"].alignment = right
    money(ws[f"F{sum_row}"])

    ws[f"E{sum_row+1}"].value = "Freight"
    ws[f"E{sum_row+1}"].alignment = right
    ws[f"F{sum_row+1}"].value = freight
    ws[f"F{sum_row+1}"].alignment = right
    money(ws[f"F{sum_row+1}"])

    ws[f"E{sum_row+2}"].value = "Total"
    ws[f"E{sum_row+2}"].alignment = right
    ws[f"E{sum_row+2}"].font = bold
    ws[f"F{sum_row+2}"].value = total_before_ppn
    ws[f"F{sum_row+2}"].alignment = right
    money(ws[f"F{sum_row+2}"])

    ws[f"E{sum_row+3}"].value = f"PPN {int(ppn_rate*100)}%"
    ws[f"E{sum_row+3}"].alignment = right
    ws[f"F{sum_row+3}"].value = ppn
    ws[f"F{sum_row+3}"].alignment = right
    money(ws[f"F{sum_row+3}"])

    ws[f"E{sum_row+4}"].value = "Less: Deposit"
    ws[f"E{sum_row+4}"].alignment = right
    ws[f"F{sum_row+4}"].value = deposit
    ws[f"F{sum_row+4}"].alignment = right
    money(ws[f"F{sum_row+4}"])

    ws[f"E{sum_row+5}"].value = "Balance Due"
    ws[f"E{sum_row+5}"].alignment = right
    ws[f"E{sum_row+5}"].font = bold
    ws[f"F{sum_row+5}"].value = balance
    ws[f"F{sum_row+5}"].alignment = right
    ws[f"F{sum_row+5}"].font = bold
    money(ws[f"F{sum_row+5}"])

    _set_border(ws, sum_row, 5, sum_row + 5, 6, border)

    pay_row = sum_row
    ws.merge_cells(f"A{pay_row}:D{pay_row}")
    ws[f"A{pay_row}"].value = "Please Transfer Full Amount to:"
    ws[f"A{pay_row}"].font = bold
    ws[f"A{pay_row}"].alignment = left

    ws[f"A{pay_row+1}"].value = "Beneficiary :"
    ws.merge_cells(f"B{pay_row+1}:D{pay_row+1}")
    ws[f"B{pay_row+1}"].value = payment["beneficiary"]

    ws[f"A{pay_row+2}"].value = "Bank Name :"
    ws.merge_cells(f"B{pay_row+2}:D{pay_row+2}")
    ws[f"B{pay_row+2}"].value = payment["bank_name"]

    ws[f"A{pay_row+3}"].value = "Branch :"
    ws.merge_cells(f"B{pay_row+3}:D{pay_row+3}")
    ws[f"B{pay_row+3}"].value = payment["branch"]

    ws[f"A{pay_row+4}"].value = "IDR Acct :"
    ws.merge_cells(f"B{pay_row+4}:D{pay_row+4}")
    ws[f"B{pay_row+4}"].value = payment["idr_acct"]

    _set_border(ws, pay_row, 1, pay_row + 4, 4, border)

    try:
        folder = str(FILES_DIR)
    except Exception:
        folder = "static/files"
    os.makedirs(folder, exist_ok=True)

    out_path = os.path.join(folder, f"{fname_base}.xlsx")
    wb.save(out_path)
    return f"{fname_base}.xlsx"

def create_invoice_pdf(inv: dict, fname_base: str) -> str:
    try:
        folder = str(FILES_DIR)
    except Exception:
        folder = "static/files"
    os.makedirs(folder, exist_ok=True)

    out_path = os.path.join(folder, f"{fname_base}.pdf")
    c = canvas.Canvas(out_path, pagesize=A4)
    width, height = A4

    def draw_text(x, y, txt, size=10, bold=False):
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        c.drawString(x, y, txt)

    invoice_no = inv.get("invoice_no") or ""
    inv_date = inv.get("invoice_date") or ""
    bill_to = inv.get("bill_to") or {}
    ship_to = inv.get("ship_to") or {}
    phone = inv.get("phone") or ""
    fax = inv.get("fax") or ""
    attn = inv.get("attn") or ""
    sales_person = inv.get("sales_person") or "Syaeful Bakri"
    ref_no = inv.get("ref_no") or ""
    ship_via = inv.get("ship_via") or ""
    ship_date = inv.get("ship_date") or ""
    terms = inv.get("terms") or ""
    no_surat_jalan = inv.get("no_surat_jalan") or ""
    items = inv.get("items") or []
    freight = int(inv.get("freight") or 0)
    ppn_rate = float(inv.get("ppn_rate") or 0.11)
    deposit = int(inv.get("deposit") or 0)
    payment = inv.get("payment") or {}

    y = height - 50
    draw_text(40, y, "INVOICE", 18, True)
    draw_text(380, y, f"Invoice: {invoice_no}", 10, True)
    y -= 14
    draw_text(380, y, f"Date: {inv_date}", 10, True)
    y -= 14
    draw_text(380, y, f"No. Surat Jalan: {no_surat_jalan}", 9, False)

    y -= 30
    draw_text(40, y, "Bill To:", 11, True)
    draw_text(300, y, "Ship To:", 11, True)

    y -= 14
    bt_lines = [bill_to.get("name",""), bill_to.get("address",""), bill_to.get("address2","")]
    st_lines = [ship_to.get("name",""), ship_to.get("address",""), ship_to.get("address2","")]
    bt_lines = [x for x in bt_lines if (x or "").strip()]
    st_lines = [x for x in st_lines if (x or "").strip()]

    yy = y
    for line in bt_lines[:4]:
        draw_text(40, yy, str(line), 9, False)
        yy -= 12
    yy2 = y
    for line in st_lines[:4]:
        draw_text(300, yy2, str(line), 9, False)
        yy2 -= 12

    y = min(yy, yy2) - 10
    draw_text(40, y, f"Phone: {phone}", 9, False)
    draw_text(300, y, f"Fax: {fax}", 9, False)
    y -= 12
    draw_text(40, y, f"Attn: {attn}", 9, False)

    y -= 18
    draw_text(40, y, f"Ref No.: {ref_no}", 9, False)
    draw_text(170, y, f"Sales Person: {sales_person}", 9, False)
    draw_text(380, y, f"Ship Via: {ship_via}", 9, False)
    y -= 12
    draw_text(380, y, f"Ship Date: {ship_date}", 9, False)
    draw_text(40, y, f"Terms: {terms}", 9, False)

    y -= 20
    table_top = y
    table_height = 220
    c.rect(40, table_top - table_height, width - 80, table_height, stroke=1, fill=0)

    header_y = table_top - 18
    draw_text(50, header_y, "Qty", 9, True)
    draw_text(90, header_y, "Unit", 9, True)
    draw_text(130, header_y, "Date", 9, True)
    draw_text(190, header_y, "Description", 9, True)
    draw_text(430, header_y, "Price", 9, True)
    draw_text(500, header_y, "Amount", 9, True)
    c.line(40, header_y - 6, width - 40, header_y - 6)

    row_y = header_y - 20
    subtotal = 0
    max_rows = 10
    for idx in range(max_rows):
        if idx < len(items):
            it = items[idx]
            qty = it.get("qty") or 0
            unit = it.get("unit") or ""
            dt = it.get("date") or inv_date
            desc = it.get("description") or ""
            price = int(it.get("price") or 0)
            amount = int(round(float(qty) * price))
            subtotal += amount

            draw_text(50, row_y, str(qty), 9, False)
            draw_text(90, row_y, str(unit), 9, False)
            draw_text(130, row_y, str(dt), 9, False)
            draw_text(190, row_y, str(desc)[:55], 9, False)
            draw_text(430, row_y, f"{price:,}".replace(",", "."), 9, False)
            draw_text(500, row_y, f"{amount:,}".replace(",", "."), 9, False)
        row_y -= 18

    total_before_ppn = subtotal + freight
    ppn = int(round(total_before_ppn * ppn_rate))
    balance = total_before_ppn + ppn - deposit

    y2 = table_top - table_height - 20
    draw_text(380, y2, "Total:", 10, True)
    draw_text(500, y2, f"{subtotal:,}".replace(",", "."), 10, False)
    y2 -= 14
    draw_text(380, y2, "Freight:", 10, True)
    draw_text(500, y2, f"{freight:,}".replace(",", "."), 10, False)
    y2 -= 14
    draw_text(380, y2, "Total:", 10, True)
    draw_text(500, y2, f"{total_before_ppn:,}".replace(",", "."), 10, False)
    y2 -= 14
    draw_text(380, y2, f"PPN {int(ppn_rate*100)}%:", 10, True)
    draw_text(500, y2, f"{ppn:,}".replace(",", "."), 10, False)
    y2 -= 14
    draw_text(380, y2, "Less: Deposit:", 10, True)
    draw_text(500, y2, f"{deposit:,}".replace(",", "."), 10, False)
    y2 -= 16
    draw_text(380, y2, "Balance Due:", 11, True)
    draw_text(500, y2, f"{balance:,}".replace(",", "."), 11, True)

    y3 = 90
    draw_text(40, y3 + 40, "Please Transfer Full Amount to:", 10, True)
    draw_text(40, y3 + 24, f"Beneficiary : {payment.get('beneficiary','')}", 9, False)
    draw_text(40, y3 + 12, f"Bank Name   : {payment.get('bank_name','')}", 9, False)
    draw_text(40, y3 + 0,  f"Branch      : {payment.get('branch','')}", 9, False)
    draw_text(40, y3 - 12, f"IDR Acct    : {payment.get('idr_acct','')}", 9, False)

    c.showPage()
    c.save()
    return f"{fname_base}.pdf"

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/history", methods=["GET"])
def api_history_list():
    try:
        q = (request.args.get("q") or "").strip()
        items = db_list_histories(limit=200, q=q if q else None)
        return jsonify({"items": items})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/history/<int:history_id>", methods=["GET"])
def api_history_detail(history_id):
    try:
        detail = db_get_history_detail(history_id)
        if not detail:
            return jsonify({"error": "history tidak ditemukan"}), 404

        return jsonify({
            "id": detail["id"],
            "title": detail["title"],
            "task_type": detail["task_type"],
            "created_at": detail["created_at"],
            "data": json.loads(detail.get("data_json") or "{}"),
            "files": json.loads(detail.get("files_json") or "[]"),
            "messages": json.loads(detail.get("messages_json") or "[]"),
            "state": json.loads(detail.get("state_json") or "{}"),
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/history/<int:history_id>", methods=["PUT"])
def api_history_update(history_id):
    try:
        body = request.get_json() or {}
        new_title = (body.get("title") or "").strip()
        if not new_title:
            return jsonify({"error": "title wajib diisi"}), 400
        ok = db_update_title(history_id, new_title)
        if not ok:
            return jsonify({"error": "history tidak ditemukan"}), 404
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/history/<int:history_id>", methods=["DELETE"])
def api_history_delete(history_id):
    try:
        ok = db_delete_history(history_id)
        if not ok:
            return jsonify({"error": "history tidak ditemukan"}), 404
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/documents", methods=["GET"])
def api_documents():
    try:
        q = (request.args.get("q") or "").strip().lower()
        items = db_list_histories(limit=500)

        docs = []
        for h in items:
            detail = db_get_history_detail(int(h["id"]))
            if not detail:
                continue
            try:
                files = json.loads(detail.get("files_json") or "[]")
            except:
                files = []
            for f in files:
                filename = (f.get("filename") or "").strip()
                if not filename:
                    continue
                title = detail.get("title") or ""
                task_type = detail.get("task_type") or ""
                created_at = detail.get("created_at") or ""

                row = {
                    "history_id": int(detail["id"]),
                    "history_title": title,
                    "task_type": task_type,
                    "created_at": created_at,
                    "type": f.get("type"),
                    "filename": filename,
                    "url": f.get("url"),
                }

                if q:
                    hay = f"{title} {filename} {task_type}".lower()
                    if q not in hay:
                        continue

                docs.append(row)

        docs.sort(key=lambda x: x.get("created_at") or "", reverse=True)
        return jsonify({"items": docs})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/chat", methods=["POST"])
def chat():
    try:
        data = request.get_json() or {}
        text = (data.get("message", "") or "").strip()
        history_id_in = data.get("history_id")

        if not text:
            return jsonify({"error": "Pesan kosong"}), 400

        sid = request.headers.get("X-Session-ID") or session.get("sid")
        if not sid:
            sid = str(uuid.uuid4())
            session["sid"] = sid

        state = conversations.get(sid, {'step': 'idle', 'data': {}})
        lower = text.lower()

        if history_id_in:
            try:
                db_append_message(int(history_id_in), "user", text, files=[])
                db_update_state(int(history_id_in), state)
            except:
                pass

        if (("invoice" in lower) or ("faktur" in lower)) and (state.get("step") == "idle"):
            inv_no = get_next_invoice_no()
            state["step"] = "inv_billto_name"
            state["data"] = {
                "invoice_no": inv_no,
                "invoice_date": datetime.now().strftime("%d-%b-%y"),
                "bill_to": {"name": "", "address": "", "address2": ""},
                "ship_to": {"name": "", "address": "", "address2": ""},
                "phone": "",
                "fax": "",
                "attn": "Accounting / Finance",
                "sales_person": "Syaeful Bakri",
                "ref_no": "",
                "ship_via": "",
                "ship_date": "",
                "terms": "",
                "no_surat_jalan": "",
                "items": [],
                "current_item": {},
                "freight": 0,
                "ppn_rate": 0.11,
                "deposit": 0,
                "payment": {
                    "beneficiary": "PT. Sarana Trans Bersama Jaya",
                    "bank_name": "BCA",
                    "branch": "Cibadak - Sukabumi",
                    "idr_acct": "35212 26666",
                }
            }
            conversations[sid] = state

            out_text = (
                "Baik, saya bantu buatkan <b>INVOICE</b>.<br><br>"
                f"✅ Invoice No: <b>{inv_no}</b><br>"
                f"✅ Date: <b>{state['data']['invoice_date']}</b><br><br>"
                "❓ <b>1. Bill To - Nama Perusahaan?</b>"
            )

            history_id_created = None
            if not history_id_in:
                history_id_created = db_insert_history(
                    title="Chat Baru",
                    task_type=data.get("taskType") or "invoice",
                    data={},
                    files=[],
                    messages=[
                        {"id": uuid.uuid4().hex[:12], "sender": "user", "text": text, "files": [], "timestamp": datetime.now().isoformat()},
                        {"id": uuid.uuid4().hex[:12], "sender": "assistant", "text": re.sub(r'<br\s*/?>', '\n', out_text), "files": [], "timestamp": datetime.now().isoformat()},
                    ],
                    state=state
                )
            else:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)

            return jsonify({"text": out_text, "history_id": history_id_created or history_id_in})

        if state.get("step") == "inv_billto_name":
            state["data"]["bill_to"]["name"] = text.strip()

            alamat = search_company_address(text).strip()
            if not alamat:
                alamat = search_company_address_ai(text).strip()
            if not alamat:
                alamat = "Di Tempat"

            state["data"]["bill_to"]["address"] = alamat
            state["step"] = "inv_shipto_same"
            conversations[sid] = state

            out_text = (
                f"✅ Bill To: <b>{state['data']['bill_to']['name']}</b><br>"
                f"✅ Alamat: <b>{alamat}</b><br><br>"
                "❓ <b>2. Ship To sama dengan Bill To?</b> (ya/tidak)"
            )

            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get("step") == "inv_shipto_same":
            if ("ya" in lower) or ("iya" in lower):
                state["data"]["ship_to"] = dict(state["data"]["bill_to"])
                state["step"] = "inv_phone"
                conversations[sid] = state
                out_text = "❓ <b>3. Phone?</b> (boleh kosong, ketik '-' jika tidak ada)"
            elif ("tidak" in lower) or ("gak" in lower) or ("nggak" in lower):
                state["step"] = "inv_shipto_name"
                conversations[sid] = state
                out_text = "❓ <b>2A. Ship To - Nama Perusahaan?</b>"
            else:
                out_text = "⚠️ Mohon jawab dengan <b>'ya'</b> atau <b>'tidak'</b><br><br>❓ <b>2. Ship To sama dengan Bill To?</b>"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get("step") == "inv_shipto_name":
            state["data"]["ship_to"]["name"] = text.strip()

            alamat = search_company_address(text).strip()
            if not alamat:
                alamat = search_company_address_ai(text).strip()
            if not alamat:
                alamat = "Di Tempat"
            state["data"]["ship_to"]["address"] = alamat

            state["step"] = "inv_phone"
            conversations[sid] = state
            out_text = "❓ <b>3. Phone?</b> (boleh kosong, ketik '-' jika tidak ada)"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get("step") == "inv_phone":
            state["data"]["phone"] = "" if text.strip() in ("-", "") else text.strip()
            state["step"] = "inv_fax"
            conversations[sid] = state
            out_text = "❓ <b>4. Fax?</b> (boleh kosong, ketik '-' jika tidak ada)"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get("step") == "inv_fax":
            state["data"]["fax"] = "" if text.strip() in ("-", "") else text.strip()
            state["step"] = "inv_attn"
            conversations[sid] = state
            out_text = "❓ <b>5. Attn?</b> (default: Accounting / Finance | ketik '-' untuk default)"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get("step") == "inv_attn":
            if text.strip() not in ("-", ""):
                state["data"]["attn"] = text.strip()
            state["step"] = "inv_item_qty"
            state["data"]["current_item"] = {}
            conversations[sid] = state
            out_text = "📦 <b>Item #1</b><br>❓ <b>6. Qty?</b> (contoh: 749 atau 3,5)"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get("step") == "inv_item_qty":
            qty = parse_qty_id(text)
            state["data"]["current_item"]["qty"] = qty
            state["step"] = "inv_item_unit"
            conversations[sid] = state
            out_text = "❓ <b>6A. Unit?</b> (contoh: Kg / Liter / Pcs)"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get("step") == "inv_item_unit":
            state["data"]["current_item"]["unit"] = text.strip()
            state["data"]["current_item"]["date"] = state["data"]["invoice_date"]
            state["step"] = "inv_item_desc"
            conversations[sid] = state
            out_text = "❓ <b>6B. Jenis Limbah / Kode Limbah?</b><br><i>(Contoh: 'A102d' atau 'aki baterai bekas' | atau ketik <b>NON B3</b>)</i>"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get("step") == "inv_item_desc":
            if is_non_b3_input(text):
                state["data"]["current_item"]["description"] = ""
                state["step"] = "inv_item_desc_manual"
                conversations[sid] = state
                out_text = "❓ <b>6C. Deskripsi (manual) apa?</b>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})

            kode, data_limbah = find_limbah_by_kode(text)
            if not (kode and data_limbah):
                kode, data_limbah = find_limbah_by_jenis(text)

            if kode and data_limbah:
                state["data"]["current_item"]["description"] = data_limbah["jenis"]
                state["step"] = "inv_item_price"
                conversations[sid] = state
                out_text = f"✅ Deskripsi: <b>{data_limbah['jenis']}</b><br><br>❓ <b>6D. Price (Rp)?</b>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})

            out_text = f"❌ Maaf, limbah '<b>{text}</b>' tidak ditemukan.<br><br>Ketik kode/jenis lain atau <b>NON B3</b>."
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get("step") == "inv_item_desc_manual":
            state["data"]["current_item"]["description"] = text.strip()
            state["step"] = "inv_item_price"
            conversations[sid] = state
            out_text = "❓ <b>6D. Price (Rp)?</b>"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get("step") == "inv_item_price":
            price = parse_amount_id(text)
            state["data"]["current_item"]["price"] = price
            state["data"]["items"].append(state["data"]["current_item"])
            state["data"]["current_item"] = {}
            state["step"] = "inv_add_more_item"
            conversations[sid] = state
            out_text = "❓ <b>Tambah item lagi?</b> (ya/tidak)"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get("step") == "inv_add_more_item":
            if ("ya" in lower) or ("iya" in lower):
                num = len(state["data"]["items"])
                state["step"] = "inv_item_qty"
                state["data"]["current_item"] = {}
                conversations[sid] = state
                out_text = f"📦 <b>Item #{num+1}</b><br>❓ <b>6. Qty?</b>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})

            if ("tidak" in lower) or ("gak" in lower) or ("nggak" in lower) or ("skip" in lower) or ("lewat" in lower):
                state["step"] = "inv_freight"
                conversations[sid] = state
                out_text = "❓ <b>7. Biaya Transportasi/Freight (Rp)?</b> (0 jika tidak ada)"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})

            out_text = "⚠️ Mohon jawab <b>ya</b> atau <b>tidak</b>."
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get("step") == "inv_freight":
            state["data"]["freight"] = parse_amount_id(text)
            state["step"] = "inv_deposit"
            conversations[sid] = state
            out_text = "❓ <b>8. Deposit (Rp)?</b> (0 jika tidak ada)"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get("step") == "inv_deposit":
            state["data"]["deposit"] = parse_amount_id(text)

            nama_pt_raw = (state["data"].get("bill_to") or {}).get("name", "").strip()
            safe_pt = re.sub(r'[^A-Za-z0-9 \-]+', '', nama_pt_raw).strip()
            safe_pt = re.sub(r'\s+', ' ', safe_pt).strip()
            base_fname = f"Invoice - {safe_pt}" if safe_pt else "Invoice"
            fname_base = make_unique_filename_base(base_fname)

            xlsx = create_invoice_xlsx(state["data"], fname_base)
            pdf_preview = create_invoice_pdf(state["data"], fname_base)

            files = [
                {"type": "xlsx", "filename": xlsx, "url": f"/download/{xlsx}"},
                {"type": "pdf", "filename": pdf_preview, "url": f"/download/{pdf_preview}"},
            ]

            conversations[sid] = {'step': 'idle', 'data': {}}

            history_title = f"Invoice {nama_pt_raw}" if nama_pt_raw else "Invoice"
            history_task_type = "invoice"

            if history_id_in:
                from utils import db_connect
                conn = db_connect()
                cur = conn.cursor()
                cur.execute("""
                    UPDATE chat_history
                    SET title = ?, task_type = ?, data_json = ?, files_json = ?
                    WHERE id = ?
                """, (
                    history_title,
                    history_task_type,
                    json.dumps(state["data"], ensure_ascii=False),
                    json.dumps(files, ensure_ascii=False),
                    int(history_id_in),
                ))
                conn.commit()
                conn.close()
                history_id = int(history_id_in)
            else:
                history_id = db_insert_history(
                    title=history_title,
                    task_type=history_task_type,
                    data=state["data"],
                    files=files,
                    messages=[],
                    state={}
                )

            out_text = (
                "🎉 <b>Invoice berhasil dibuat!</b><br><br>"
                f"✅ Invoice No: <b>{state['data'].get('invoice_no')}</b><br>"
                f"✅ Bill To: <b>{(state['data'].get('bill_to') or {}).get('name','')}</b><br>"
                f"✅ Total Item: <b>{len(state['data'].get('items') or [])}</b><br><br>"
                "📌 Preview: PDF<br>"
                "📌 Download: Excel (.xlsx)"
            )

            db_append_message(history_id, "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=files)

            return jsonify({"text": out_text, "files": files, "history_id": history_id})

        if ('mou' in lower) and (state.get('step') == 'idle'):
            nomor_depan = get_next_mou_no_depan()
            state['step'] = 'mou_pihak_pertama'
            state['data'] = {
                'nomor_depan': nomor_depan,
                'nomor_surat': "",
                'items_limbah': [],
                'current_item': {},
                'pihak_kedua': "PT Sarana Trans Bersama Jaya",
                'pihak_kedua_kode': "STBJ",
                'pihak_pertama': "",
                'alamat_pihak_pertama': "",
                'pihak_ketiga': "",
                'pihak_ketiga_kode': "",
                'alamat_pihak_ketiga': "",
                'ttd_pihak_pertama': "",
                'jabatan_pihak_pertama': "",
                'ttd_pihak_ketiga': "",
                'jabatan_pihak_ketiga': "",
            }
            conversations[sid] = state

            out_text = (
                "Baik, saya bantu buatkan <b>MoU Tripartit</b>.<br><br>"
                f"✅ No Depan: <b>{nomor_depan}</b><br>"
                "✅ Nomor lengkap otomatis mengikuti format template.<br>"
                "✅ Tanggal otomatis hari ini.<br><br>"
                "❓ <b>1. Nama Perusahaan (PIHAK PERTAMA / Penghasil Limbah)?</b>"
            )

            history_id_created = None
            if not history_id_in:
                history_id_created = db_insert_history(
                    title="Chat Baru",
                    task_type=data.get("taskType") or "mou",
                    data={},
                    files=[],
                    messages=[
                        {"id": uuid.uuid4().hex[:12], "sender": "user", "text": text, "files": [], "timestamp": datetime.now().isoformat()},
                        {"id": uuid.uuid4().hex[:12], "sender": "assistant", "text": re.sub(r'<br\s*/?>', '\n', out_text), "files": [], "timestamp": datetime.now().isoformat()},
                    ],
                    state=state
                )
            else:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)

            return jsonify({"text": out_text, "history_id": history_id_created or history_id_in})

        if state.get('step') == 'mou_pihak_pertama':
            state['data']['pihak_pertama'] = text.strip()

            alamat = search_company_address(text).strip()
            if not alamat:
                alamat = search_company_address_ai(text).strip()
            if not alamat:
                alamat = "Di Tempat"

            state['data']['alamat_pihak_pertama'] = alamat
            state['step'] = 'mou_pilih_pihak_ketiga'
            conversations[sid] = state

            out_text = (
                f"✅ PIHAK PERTAMA: <b>{state['data']['pihak_pertama']}</b><br>"
                f"✅ Alamat: <b>{alamat}</b><br><br>"
                "❓ <b>2. Pilih PIHAK KETIGA (Pengelola Limbah):</b><br>"
                "1. HBSP<br>2. KJL<br>3. MBI<br>4. CGA<br><br>"
                "<i>(Ketik nomor 1-4 atau ketik HBSP/KJL/MBI/CGA)</i>"
            )

            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)

            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get('step') == 'mou_pilih_pihak_ketiga':
            pilihan = text.strip().upper()
            mapping = {"1": "HBSP", "2": "KJL", "3": "MBI", "4": "CGA", "HBSP": "HBSP", "KJL": "KJL", "MBI": "MBI", "CGA": "CGA"}
            kode = mapping.get(pilihan)
            if not kode:
                out_text = "⚠️ Pilihan tidak valid.<br><br>Pilih PIHAK KETIGA:<br>1. HBSP<br>2. KJL<br>3. MBI<br>4. CGA"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                return jsonify({"text": out_text, "history_id": history_id_in})

            pihak3_nama_map = {"HBSP": "PT Harapan Baru Sejahtera Plastik", "KJL": "KJL", "MBI": "MBI", "CGA": "CGA"}
            pihak3_alamat_map = {
                "HBSP": "Jl. Karawang – Bekasi KM. 1 Bojongsari, Kec. Kedungwaringin, Kab. Bekasi – Jawa Barat",
                "KJL": "",
                "MBI": "",
                "CGA": "",
            }

            state['data']['pihak_ketiga'] = pihak3_nama_map.get(kode, kode)
            state['data']['pihak_ketiga_kode'] = kode
            state['data']['alamat_pihak_ketiga'] = pihak3_alamat_map.get(kode, "")

            state['data']['nomor_surat'] = build_mou_nomor_surat(state['data'])
            state['step'] = 'mou_jenis_kode_limbah'
            state['data']['current_item'] = {}
            conversations[sid] = state

            out_text = (
                f"✅ PIHAK KETIGA: <b>{state['data']['pihak_ketiga']}</b><br>"
                f"✅ Nomor MoU: <b>{state['data']['nomor_surat']}</b><br><br>"
                "📦 <b>Item #1</b><br>"
                "❓ <b>3. Sebutkan Jenis Limbah atau Kode Limbah</b><br>"
                "<i>(Contoh: 'A102d' atau 'aki baterai bekas' | atau ketik <b>NON B3</b>)</i>"
            )

            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)

            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get('step') == 'mou_jenis_kode_limbah':
            if is_non_b3_input(text):
                state['data']['current_item']['kode_limbah'] = "NON B3"
                state['data']['current_item']['jenis_limbah'] = ""
                state['step'] = 'mou_manual_jenis_limbah'
                conversations[sid] = state
                out_text = "✅ Kode: <b>NON B3</b><br><br>❓ <b>3A. Jenis Limbah (manual) apa?</b>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})

            kode, data_limbah = find_limbah_by_kode(text)
            if not (kode and data_limbah):
                kode, data_limbah = find_limbah_by_jenis(text)

            if kode and data_limbah:
                state['data']['current_item']['kode_limbah'] = kode
                state['data']['current_item']['jenis_limbah'] = data_limbah['jenis']
                state['data']['items_limbah'].append(state['data']['current_item'])
                num = len(state['data']['items_limbah'])
                state['step'] = 'mou_tambah_item'
                state['data']['current_item'] = {}
                conversations[sid] = state
                out_text = f"✅ Item #{num} tersimpan!<br>• Jenis: <b>{data_limbah['jenis']}</b><br>• Kode: <b>{kode}</b><br><br>❓ <b>Tambah item lagi?</b> (ya/tidak)"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})

            out_text = f"❌ Maaf, limbah '<b>{text}</b>' tidak ditemukan.<br><br>Ketik kode/jenis lain atau <b>NON B3</b>."
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get('step') == 'mou_manual_jenis_limbah':
            state['data']['current_item']['jenis_limbah'] = text.strip()
            state['data']['items_limbah'].append(state['data']['current_item'])
            num = len(state['data']['items_limbah'])
            state['step'] = 'mou_tambah_item'
            state['data']['current_item'] = {}
            conversations[sid] = state
            out_text = f"✅ Item #{num} tersimpan!<br>• Jenis (manual): <b>{state['data']['items_limbah'][-1]['jenis_limbah']}</b><br>• Kode: <b>NON B3</b><br><br>❓ <b>Tambah item lagi?</b> (ya/tidak)"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get('step') == 'mou_tambah_item':
            if re.match(r'^\d+', text.strip()):
                out_text = "⚠️ Mohon jawab dengan <b>'ya'</b> atau <b>'tidak'</b><br><br>❓ <b>Tambah item lagi?</b>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                return jsonify({"text": out_text, "history_id": history_id_in})

            if ('ya' in lower) or ('iya' in lower):
                num = len(state['data']['items_limbah'])
                state['step'] = 'mou_jenis_kode_limbah'
                state['data']['current_item'] = {}
                conversations[sid] = state
                out_text = f"📦 <b>Item #{num+1}</b><br>❓ <b>3. Sebutkan Jenis Limbah atau Kode Limbah</b><br><i>(Contoh: 'A102d' | atau <b>NON B3</b>)</i>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})

            if ('tidak' in lower) or ('skip' in lower) or ('lewat' in lower) or ('gak' in lower) or ('nggak' in lower):
                state['step'] = 'mou_ttd_pihak_pertama'
                conversations[sid] = state
                out_text = "❓ <b>Nama penandatangan PIHAK PERTAMA?</b>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})

            out_text = "⚠️ Mohon jawab dengan <b>'ya'</b> atau <b>'tidak'</b><br><br>❓ <b>Tambah item lagi?</b>"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get('step') == 'mou_ttd_pihak_pertama':
            state['data']['ttd_pihak_pertama'] = text.strip()
            state['step'] = 'mou_jabatan_pihak_pertama'
            conversations[sid] = state
            out_text = "❓ <b>Jabatan penandatangan PIHAK PERTAMA?</b>"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get('step') == 'mou_jabatan_pihak_pertama':
            state['data']['jabatan_pihak_pertama'] = text.strip()
            state['step'] = 'mou_ttd_pihak_ketiga'
            conversations[sid] = state
            out_text = "❓ <b>Nama penandatangan PIHAK KETIGA?</b>"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get('step') == 'mou_ttd_pihak_ketiga':
            state['data']['ttd_pihak_ketiga'] = text.strip()
            state['step'] = 'mou_jabatan_pihak_ketiga'
            conversations[sid] = state
            out_text = "❓ <b>Jabatan penandatangan PIHAK KETIGA?</b>"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        if state.get('step') == 'mou_jabatan_pihak_ketiga':
            state['data']['jabatan_pihak_ketiga'] = text.strip()

            nama_pt_raw = state['data'].get('pihak_pertama', '').strip()
            safe_pt = re.sub(r'[^A-Za-z0-9 \-]+', '', nama_pt_raw).strip()
            safe_pt = re.sub(r'\s+', ' ', safe_pt).strip()

            base_fname = f"MoU - {safe_pt}" if safe_pt else "MoU - Perusahaan"
            fname_base = make_unique_filename_base(base_fname)

            if not state['data'].get("nomor_surat"):
                state['data']['nomor_surat'] = build_mou_nomor_surat(state['data'])

            docx = create_mou_docx(state['data'], fname_base)
            pdf = create_pdf(fname_base)

            conversations[sid] = {'step': 'idle', 'data': {}}

            files = [{"type": "docx", "filename": docx, "url": f"/download/{docx}"}]
            if pdf:
                files.append({"type": "pdf", "filename": pdf, "url": f"/download/{pdf}"})

            history_title = f"MoU {nama_pt_raw}" if nama_pt_raw else "MoU"
            history_task_type = "mou"

            if history_id_in:
                from utils import db_connect
                conn = db_connect()
                cur = conn.cursor()
                cur.execute("""
                    UPDATE chat_history
                    SET title = ?, task_type = ?, data_json = ?, files_json = ?
                    WHERE id = ?
                """, (
                    history_title,
                    history_task_type,
                    json.dumps(state['data'], ensure_ascii=False),
                    json.dumps(files, ensure_ascii=False),
                    int(history_id_in),
                ))
                conn.commit()
                conn.close()
                history_id = int(history_id_in)
            else:
                history_id = db_insert_history(
                    title=history_title,
                    task_type=history_task_type,
                    data=state['data'],
                    files=files,
                    messages=[],
                    state={}
                )

            out_text = (
                "🎉 <b>MoU berhasil dibuat!</b><br><br>"
                f"✅ Nomor MoU: <b>{state['data'].get('nomor_surat')}</b><br>"
                f"✅ PIHAK PERTAMA: <b>{state['data'].get('pihak_pertama')}</b><br>"
                f"✅ PIHAK KEDUA: <b>{state['data'].get('pihak_kedua')}</b><br>"
                f"✅ PIHAK KETIGA: <b>{state['data'].get('pihak_ketiga')}</b><br>"
                f"✅ Total Limbah: <b>{len(state['data'].get('items_limbah') or [])} item</b>"
            )

            db_append_message(history_id, "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=files)
            return jsonify({"text": out_text, "files": files, "history_id": history_id})

        if ('quotation' in lower or 'penawaran' in lower or ('buat' in lower and 'mou' not in lower)):
            nomor_depan = get_next_nomor()
            state['step'] = 'nama_perusahaan'
            now = datetime.now()
            state['data'] = {'nomor_depan': nomor_depan, 'items_limbah': [], 'bulan_romawi': now.strftime('%m')}
            conversations[sid] = state

            out_text = f"Baik, saya bantu buatkan quotation.<br><br>✅ Nomor Surat: <b>{nomor_depan}</b><br><br>❓ <b>1. Nama Perusahaan?</b>"

            history_id_created = None
            if not history_id_in:
                history_id_created = db_insert_history(
                    title="Chat Baru",
                    task_type=data.get("taskType") or "penawaran",
                    data={},
                    files=[],
                    messages=[
                        {"id": uuid.uuid4().hex[:12], "sender": "user", "text": text, "files": [], "timestamp": datetime.now().isoformat()},
                        {"id": uuid.uuid4().hex[:12], "sender": "assistant", "text": re.sub(r'<br\s*/?>', '\n', out_text), "files": [], "timestamp": datetime.now().isoformat()},
                    ],
                    state=state
                )
            else:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_created or history_id_in})

        if state['step'] == 'nama_perusahaan':
            state['data']['nama_perusahaan'] = text

            alamat = search_company_address(text).strip()
            if not alamat:
                alamat = search_company_address_ai(text).strip()
            if not alamat:
                alamat = "Di Tempat"

            state['data']['alamat_perusahaan'] = alamat
            state['step'] = 'jenis_kode_limbah'
            state['data']['current_item'] = {}
            conversations[sid] = state

            out_text = (
                f"✅ Nama: <b>{text}</b><br>✅ Alamat: <b>{alamat}</b><br><br>"
                f"📦 <b>Item #1</b><br>❓ <b>2. Sebutkan Jenis Limbah atau Kode Limbah</b><br>"
                f"<i>(Contoh: 'A102d' atau 'aki baterai bekas')</i>"
            )

            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        elif state['step'] == 'jenis_kode_limbah':
            if is_non_b3_input(text):
                state['data']['current_item']['kode_limbah'] = "NON B3"
                state['data']['current_item']['jenis_limbah'] = ""
                state['data']['current_item']['satuan'] = ""
                state['step'] = 'manual_jenis_limbah'
                conversations[sid] = state

                out_text = "✅ Kode: <b>NON B3</b><br><br>❓ <b>2A. Jenis Limbah (manual) apa?</b>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})

            kode, data_limbah = find_limbah_by_kode(text)
            if kode and data_limbah:
                state['data']['current_item']['kode_limbah'] = kode
                state['data']['current_item']['jenis_limbah'] = data_limbah['jenis']
                state['data']['current_item']['satuan'] = data_limbah['satuan']
                state['step'] = 'harga'
                conversations[sid] = state
                out_text = f"✅ Kode: <b>{kode}</b><br>✅ Jenis: <b>{data_limbah['jenis']}</b><br>✅ Satuan: <b>{data_limbah['satuan']}</b><br><br>❓ <b>3. Harga (Rp)?</b>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})
            else:
                kode, data_limbah = find_limbah_by_jenis(text)
                if kode and data_limbah:
                    state['data']['current_item']['kode_limbah'] = kode
                    state['data']['current_item']['jenis_limbah'] = data_limbah['jenis']
                    state['data']['current_item']['satuan'] = data_limbah['satuan']
                    state['step'] = 'harga'
                    conversations[sid] = state
                    out_text = f"✅ Kode: <b>{kode}</b><br>✅ Jenis: <b>{data_limbah['jenis']}</b><br>✅ Satuan: <b>{data_limbah['satuan']}</b><br><br>❓ <b>3. Harga (Rp)?</b>"
                    if history_id_in:
                        db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                        db_update_state(int(history_id_in), state)
                    return jsonify({"text": out_text, "history_id": history_id_in})
                else:
                    out_text = (
                        f"❌ Maaf, limbah '<b>{text}</b>' tidak ditemukan dalam database.<br><br>"
                        "Silakan coba lagi dengan:<br>"
                        "• Kode limbah (contoh: A102d, B105d)<br>"
                        "• Nama jenis limbah (contoh: aki baterai bekas, minyak pelumas bekas)<br>"
                        "• Atau ketik <b>NON B3</b> untuk input manual"
                    )
                    if history_id_in:
                        db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                        db_update_state(int(history_id_in), state)
                    return jsonify({"text": out_text, "history_id": history_id_in})

        elif state['step'] == 'manual_jenis_limbah':
            state['data']['current_item']['jenis_limbah'] = text
            state['step'] = 'manual_satuan'
            conversations[sid] = state
            out_text = f"✅ Jenis (manual): <b>{text}</b><br><br>❓ <b>2B. Satuan (manual) apa?</b>"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        elif state['step'] == 'manual_satuan':
            state['data']['current_item']['satuan'] = text
            state['step'] = 'harga'
            conversations[sid] = state
            out_text = f"✅ Satuan (manual): <b>{text}</b><br><br>❓ <b>3. Harga (Rp)?</b>"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        elif state['step'] == 'harga':
            harga_converted = parse_amount_id(text)
            state['data']['current_item']['harga'] = harga_converted
            state['data']['items_limbah'].append(state['data']['current_item'])
            num = len(state['data']['items_limbah'])
            state['step'] = 'tambah_item'
            conversations[sid] = state
            harga_formatted = format_rupiah(harga_converted)
            out_text = f"✅ Item #{num} tersimpan!<br>💰 Harga: <b>Rp {harga_formatted}</b><br><br>❓ <b>Tambah item lagi?</b> (ya/tidak)"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        elif state['step'] == 'tambah_item':
            if re.match(r'^\d+', text.strip()):
                out_text = "⚠️ Mohon jawab dengan <b>'ya'</b> atau <b>'tidak'</b><br><br>❓ <b>Tambah item lagi?</b>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                return jsonify({"text": out_text, "history_id": history_id_in})

            if 'ya' in lower or 'iya' in lower:
                num = len(state['data']['items_limbah'])
                state['step'] = 'jenis_kode_limbah'
                state['data']['current_item'] = {}
                conversations[sid] = state
                out_text = f"📦 <b>Item #{num+1}</b><br>❓ <b>2. Sebutkan Jenis Limbah atau Kode Limbah</b><br><i>(Contoh: 'A102d' atau 'aki baterai bekas')</i>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})
            elif 'tidak' in lower or 'skip' in lower or 'lewat' in lower or 'gak' in lower or 'nggak' in lower:
                state['step'] = 'harga_transportasi'
                conversations[sid] = state
                out_text = f"✅ Total: <b>{len(state['data']['items_limbah'])} item</b><br><br>❓ <b>4. Biaya Transportasi (Rp)?</b><br><i>Satuan: ritase</i>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})
            else:
                out_text = "⚠️ Mohon jawab dengan <b>'ya'</b> atau <b>'tidak'</b><br><br>❓ <b>Tambah item lagi?</b>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                return jsonify({"text": out_text, "history_id": history_id_in})

        elif state['step'] == 'harga_transportasi':
            transportasi_converted = parse_amount_id(text)
            state['data']['harga_transportasi'] = transportasi_converted
            state['step'] = 'tanya_mou'
            conversations[sid] = state
            transportasi_formatted = format_rupiah(transportasi_converted)
            out_text = f"✅ Transportasi: <b>Rp {transportasi_formatted}/ritase</b><br><br>❓ <b>5. Tambah Biaya MoU?</b> (ya/tidak)"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        elif state['step'] == 'tanya_mou':
            if re.match(r'^\d+', text.strip()):
                out_text = "⚠️ Mohon jawab dengan <b>'ya'</b> atau <b>'tidak'</b><br><br>❓ <b>5. Tambah Biaya MoU?</b> (ya/tidak)"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                return jsonify({"text": out_text, "history_id": history_id_in})

            if 'ya' in lower or 'iya' in lower:
                state['step'] = 'harga_mou'
                conversations[sid] = state
                out_text = "❓ <b>Biaya MoU (Rp)?</b><br><i>Satuan: Tahun</i>"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})
            elif 'tidak' in lower or 'skip' in lower or 'lewat' in lower or 'gak' in lower or 'nggak' in lower:
                state['data']['harga_mou'] = None
                state['step'] = 'tanya_termin'
                conversations[sid] = state
                out_text = "❓ <b>6. Edit Termin Pembayaran?</b><br><i>Default: 14 hari</i><br>(ketik angka atau 'tidak' untuk default)"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                    db_update_state(int(history_id_in), state)
                return jsonify({"text": out_text, "history_id": history_id_in})
            else:
                out_text = "⚠️ Mohon jawab dengan <b>'ya'</b> atau <b>'tidak'</b><br><br>❓ <b>5. Tambah Biaya MoU?</b> (ya/tidak)"
                if history_id_in:
                    db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                return jsonify({"text": out_text, "history_id": history_id_in})

        elif state['step'] == 'harga_mou':
            mou_converted = parse_amount_id(text)
            state['data']['harga_mou'] = mou_converted
            state['step'] = 'tanya_termin'
            conversations[sid] = state
            mou_formatted = format_rupiah(mou_converted)
            out_text = f"✅ MoU: <b>Rp {mou_formatted}/Tahun</b><br><br>❓ <b>6. Edit Termin Pembayaran?</b><br><i>Default: 14 hari</i><br>(ketik angka atau 'tidak' untuk default)"
            if history_id_in:
                db_append_message(int(history_id_in), "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=[])
                db_update_state(int(history_id_in), state)
            return jsonify({"text": out_text, "history_id": history_id_in})

        elif state['step'] == 'tanya_termin':
            if 'tidak' in lower or 'skip' in lower or 'lewat' in lower:
                state['data']['termin_hari'] = '14'
            else:
                state['data']['termin_hari'] = parse_termin_days(text, default=14, min_days=1, max_days=365)

            nama_pt_raw = state['data'].get('nama_perusahaan', '').strip()
            safe_pt = re.sub(r'[^A-Za-z0-9 \-]+', '', nama_pt_raw).strip()
            safe_pt = re.sub(r'\s+', ' ', safe_pt).strip()
            base_fname = f"Quotation - {safe_pt}" if safe_pt else "Quotation - Penawaran"
            fname = make_unique_filename_base(base_fname)

            docx = create_docx(state['data'], fname)
            pdf = create_pdf(fname)

            conversations[sid] = {'step': 'idle', 'data': {}}

            files = [{"type": "docx", "filename": docx, "url": f"/download/{docx}"}]
            if pdf:
                files.append({"type": "pdf", "filename": pdf, "url": f"/download/{pdf}"})

            nama_pt = state['data'].get('nama_perusahaan', '').strip()
            history_title = f"Penawaran {nama_pt}" if nama_pt else "Penawaran"
            history_task_type = "penawaran"

            if history_id_in:
                from utils import db_connect
                conn = db_connect()
                cur = conn.cursor()
                cur.execute("""
                    UPDATE chat_history
                    SET title = ?, task_type = ?, data_json = ?, files_json = ?
                    WHERE id = ?
                """, (
                    history_title,
                    history_task_type,
                    json.dumps(state['data'], ensure_ascii=False),
                    json.dumps(files, ensure_ascii=False),
                    int(history_id_in),
                ))
                conn.commit()
                conn.close()
                history_id = int(history_id_in)
            else:
                history_id = db_insert_history(
                    title=history_title,
                    task_type=history_task_type,
                    data=state['data'],
                    files=files,
                    messages=[],
                    state={}
                )

            termin_terbilang = angka_ke_terbilang(state['data']['termin_hari'])
            out_text = f"✅ Termin: <b>{state['data']['termin_hari']} ({termin_terbilang}) hari</b><br><br>🎉 <b>Quotation berhasil dibuat!</b>"
            db_append_message(history_id, "assistant", re.sub(r'<br\s*/?>', '\n', out_text), files=files)

            return jsonify({"text": out_text, "files": files, "history_id": history_id})

        ai_out = call_ai(text)
        if history_id_in:
            db_append_message(int(history_id_in), "assistant", ai_out, files=[])
        return jsonify({"text": ai_out, "history_id": history_id_in})

    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/download/<path:filename>")
def download(filename):
    return send_from_directory(str(FILES_DIR), filename, as_attachment=True)

if __name__ == "__main__":
    port = FLASK_PORT
    debug_mode = FLASK_DEBUG

    print("\n" + "="*60)
    print("🚀 QUOTATION GENERATOR")
    print("="*60)
    print(f"📁 Template: {TEMPLATE_FILE.exists() and '✅ Found' or '❌ Missing'}")
    print(f"🔑 OpenRouter: {OPENROUTER_API_KEY and '✅' or '❌'}")
    print(f"🔎 Serper: {SERPER_API_KEY and '✅' or '❌'}")
    print(f"📄 PDF: {PDF_AVAILABLE and f'✅ {PDF_METHOD}' or '❌ Disabled'}")
    print(f"🗄️  Database: {len(LIMBAH_B3_DB)} jenis limbah")
    print(f"🔢 Counter: {load_counter()}")
    print(f"🌐 Port: {port}")
    print(f"💻 Platform: {platform.system()}")
    print("="*60 + "\n")

    app.run(host="0.0.0.0", port=port, debug=debug_mode)
